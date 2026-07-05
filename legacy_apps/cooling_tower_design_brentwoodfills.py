# cooling_tower_design_complete_with_CF1200_counterflow_v2.py
# Complete Cooling Tower Design Tool with Enhanced UI & Input Controls

import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
import datetime
import hashlib
import hmac
import math
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
import io
import re
import os
import json
from pathlib import Path
import pdfplumber
import docx
from docx import Document
from docx.shared import Pt


# =============================================================================
# PASSWORD PROTECTION
# ============================================================================

def check_password():
    """Returns True if the user entered the correct password.

    Password is read from Streamlit Secrets (recommended):
      - st.secrets["APP_PASSWORD"]  (preferred)
      - st.secrets["app_password"]  (fallback)
    You can set this in Streamlit Cloud: App → Settings → Secrets.

    No password is hard-coded in the code.
    """

    # Fetch password from secrets (and optionally env var for local runs)
    expected_password = None
    try:
        if "APP_PASSWORD" in st.secrets:
            expected_password = st.secrets["APP_PASSWORD"]
        elif "app_password" in st.secrets:
            expected_password = st.secrets["app_password"]
    except Exception:
        # If secrets aren't configured or available, we'll fall back below
        expected_password = None

    if expected_password is None:
        import os
        expected_password = os.environ.get("APP_PASSWORD")

    if not expected_password:
        st.error("App password is not configured. Add APP_PASSWORD to Streamlit Secrets.")
        st.stop()

    def password_entered():
        entered = st.session_state.get("password", "")
        # constant-time comparison to avoid leaking info via timing
        if hmac.compare_digest(str(entered), str(expected_password)):
            st.session_state["password_correct"] = True
            st.session_state.pop("password", None)
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.text_input("Enter Password", type="password", on_change=password_entered, key="password")
        return False
    elif not st.session_state["password_correct"]:
        st.text_input("Enter Password", type="password", on_change=password_entered, key="password")
        st.error("😕 Password incorrect")
        return False
    else:
        return True

# ============================================================================
# ENHANCED BRENTWOOD FILL DATABASE - ADDED CF1200 AND CT1200AT
# ============================================================================

BRENTWOOD_FILLS = {
    "CF1200": {
        "name": "Brentwood ACCU-PAK CF1200 (Old Data)",
        "surface_area": 226,  # m²/m³
        "sheet_spacing": 11.7,  # mm
        "flute_angle": 30,  # degrees
        "channel_depth": 9.0,  # mm
        "channel_width": 13.5,  # mm
        "hydraulic_diameter": 8.8,  # mm
        "free_area_fraction": 0.89,
        "free_area_source": "Engineering estimate - Brentwood brochure gives sheet spacing/surface area but not net open area fraction",
        "water_passage_area": 0.78,
        "material_thickness_options": [0.20, 0.25, 0.30],
        "dry_weight_range": [36.8, 60.9],
        "water_film_thickness": 0.6,  # mm
        "max_water_loading": 14,  # m³/h·m² (lower than XF75)
        "min_water_loading": 6,
        "recommended_air_velocity": 2.2,  # m/s (lower than XF75)
        "max_air_velocity": 2.8,
        "fouling_factor": 0.80,  # Worse fouling resistance
        
        # CF1200 PERFORMANCE CURVE (TUNED TO MATCH SUPPLIER'S SAA15 DESIGN)
        # Lower efficiency than XF75 - matches supplier's Ka/L = 0.982 for L/G=2.313
        "performance_data": {
            "L_G": [0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0, 2.25, 2.5],
            "Ka_L": [1.9, 1.6, 1.35, 1.15, 0.98, 0.85, 0.75, 0.67, 0.60],  # Lower than XF75
            "delta_P_base": [45, 58, 75, 96, 122, 152, 186, 225, 268]  # Higher than XF75
        },
        
        "description": "Older cross-fluted fill with lower thermal efficiency, matches SAA15 supplier design"
    },
    
    "CT1200AT": {
        "name": "Brentwood CT1200AT (New from PDF)",
        "surface_area": 226,  # m²/m³ (same as CF1200)
        "sheet_spacing": 11.7,  # mm
        "flute_angle": 30,  # degrees
        "channel_depth": 9.0,  # mm
        "channel_width": 13.5,  # mm
        "hydraulic_diameter": 8.8,  # mm
        "free_area_fraction": 0.89,
        "free_area_source": "Engineering estimate - Brentwood brochure gives sheet spacing/surface area but not net open area fraction",
        "water_passage_area": 0.78,
        "material_thickness_options": [0.20, 0.25, 0.30],
        "dry_weight_range": [36.8, 60.9],
        "water_film_thickness": 0.6,  # mm
        "max_water_loading": 14,  # m³/h·m²
        "min_water_loading": 6,
        "recommended_air_velocity": 2.2,  # m/s
        "max_air_velocity": 2.8,
        "fouling_factor": 0.85,  # Better than CF1200
        
        # CT1200AT PERFORMANCE DATA FROM PDF
        # Based on the provided PDF graphs for different fill heights
        "performance_data": {
            "L_G": [0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0, 2.25, 2.5],
            # Ka/L values for CT1200AT (estimated from PDF performance curve)
            "Ka_L": [2.1, 1.8, 1.55, 1.35, 1.18, 1.05, 0.94, 0.85, 0.78],
            # Base pressure drop for 24" (610mm) fill height
            "delta_P_base_24": [40, 52, 68, 88, 112, 140, 172, 208, 248],
            "delta_P_base_36": [55, 72, 94, 122, 155, 194, 238, 288, 343],
            "delta_P_base_48": [70, 92, 120, 156, 198, 248, 304, 368, 438]
        },
        
        "description": "Updated CT1200AT fill with actual performance data from PDF graphs"
    },

    "CF1900SBMA": {
        "name": "Brentwood CF1900SB/MA (2017 Test Cell)",
        "surface_area": 157.5,  # m²/m³ (48 ft²/ft³)
        "sheet_spacing": 19.0,  # mm (0.75")
        "flute_angle": 31,      # degrees (CF1900)
        "channel_depth": 12.0,  # mm (approx)
        "channel_width": 18.0,  # mm (approx)
        "hydraulic_diameter": 12.5,  # mm (approx)
        "free_area_fraction": 0.86,
        "free_area_source": "Engineering estimate - Brentwood brochure gives sheet spacing/surface area but not net open area fraction",
        "water_passage_area": 0.74,
        "material_thickness_options": [0.25, 0.38],
        "dry_weight_range": [27.2, 38.4],  # kg/m³ (typical range)
        "water_film_thickness": 0.7,  # mm
        "max_water_loading": 25,  # m³/h·m² (per SI pressure drop plots)
        "min_water_loading": 10,
        "recommended_air_velocity": 2.5,  # m/s (typical)
        "max_air_velocity": 4.0,
        "fouling_factor": 0.85,

        # Placeholder Ka/L curve (to be refined if you later digitize the KaV/L graph precisely)
        "performance_data": {
            "L_G": [0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0, 2.25, 2.5],
            "Ka_L": [1.6, 1.35, 1.15, 0.98, 0.85, 0.75, 0.67, 0.60, 0.55],
            "delta_P_base": [35, 46, 60, 78, 100, 126, 155, 190, 230]
        },

        "description": "Cross-fluted film fill. Common for counterflow towers and also usable in crossflow towers."
    },

    "XF75": {
        "name": "Brentwood XF75",
        # Brentwood Herringbone Fill Product Details:
        # product surface area = 51 ft²/ft³ = 167.4 m²/m³, sheet spacing = 0.75 in = 19 mm, 16 sheets/ft.
        "surface_area": 167.4,
        "sheet_spacing": 19.0,
        "sheets_per_foot": 16,
        "flute_angle": 30,
        "channel_depth": 12.0,
        "channel_width": 18.0,
        "hydraulic_diameter": 12.5,
        "free_area_fraction": 0.91,
        "free_area_source": "Engineering estimate - not directly published in available brochure",
        "water_passage_area": 0.82,
        "material_thickness_options": [0.25, 0.38, 0.50],
        "dry_weight_range": [27.2, 52.9],
        "water_film_thickness": 0.8,
        "max_water_loading": 70,
        "min_water_loading": 20,
        "recommended_air_velocity": 2.5,
        "max_air_velocity": 4.0,
        "fouling_factor": 0.85,
        
        "performance_data": {
            "L_G": [0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0],
            "Ka_L": [2.3, 1.95, 1.65, 1.4, 1.2, 1.05, 0.92],
            "delta_P_base": [35, 45, 60, 80, 105, 135, 170]
        },
        
        "description": "High density cross-fluted fill with maximum surface area"
    },
    
    "ThermaCross": {
        "name": "Brentwood ThermaCross",
        "surface_area": 154,
        "sheet_spacing": 19.0,
        "flute_angle": 22,
        "channel_depth": 11.0,
        "channel_width": 16.5,
        "hydraulic_diameter": 10.5,
        "free_area_fraction": 0.91,
        "free_area_source": "Engineering estimate - not directly published in available brochure",
        "water_passage_area": 0.82,
        "material_thickness_options": [0.25, 0.38, 0.50],
        "dry_weight_range": [27.2, 52.9],
        "water_film_thickness": 0.8,
        "max_water_loading": 18,
        "min_water_loading": 6,
        "recommended_air_velocity": 2.4,
        "max_air_velocity": 2.9,
        "fouling_factor": 0.90,
        
        "performance_data": {
            "L_G": [0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0],
            "Ka_L": [2.0, 1.7, 1.45, 1.25, 1.08, 0.95, 0.84],
            "delta_P_base": [25, 32, 42, 55, 72, 92, 115]
        },
        
        "description": "Balanced performance fill with good thermal and hydraulic characteristics"
    },
    
    "XF125": {
        "name": "Brentwood XF125",
        "surface_area": 157.5,
        "sheet_spacing": 19.0,
        "flute_angle": 31,
        "channel_depth": 11.0,
        "channel_width": 16.5,
        "hydraulic_diameter": 10.5,
        "free_area_fraction": 0.91,
        "free_area_source": "Engineering estimate - not directly published in available brochure",
        "water_passage_area": 0.82,
        "material_thickness_options": [0.25, 0.38, 0.50],
        "dry_weight_range": [27.2, 52.9],
        "water_film_thickness": 0.8,
        "max_water_loading": 18,
        "min_water_loading": 6,
        "recommended_air_velocity": 2.4,
        "max_air_velocity": 2.9,
        "fouling_factor": 0.88,
        
        "performance_data": {
            "L_G": [0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0],
            "Ka_L": [2.1, 1.78, 1.52, 1.3, 1.12, 0.98, 0.87],
            "delta_P_base": [28, 36, 47, 62, 80, 102, 128]
        },
        
        "description": "Optimized flute angle for enhanced heat transfer with moderate pressure drop"
    },
    
    "XF3000": {
        "name": "Brentwood XF3000",
        "surface_area": 102,
        "sheet_spacing": 30.5,
        "flute_angle": 30,
        "channel_depth": 13.5,
        "channel_width": 22.5,
        "hydraulic_diameter": 14.2,
        "free_area_fraction": 0.93,
        "free_area_source": "Engineering estimate - not directly published in available brochure",
        "water_passage_area": 0.85,
        "material_thickness_options": [0.38, 0.51],
        "dry_weight_range": [25.6, 35.2],
        "water_film_thickness": 1.2,
        "max_water_loading": 25,
        "min_water_loading": 8,
        "recommended_air_velocity": 2.6,
        "max_air_velocity": 3.2,
        "fouling_factor": 0.95,
        
        "performance_data": {
            "L_G": [0.5, 0.75, 1.0, 1.25, 1.5, 1.75, 2.0],
            "Ka_L": [1.7, 1.45, 1.25, 1.08, 0.94, 0.83, 0.74],
            "delta_P_base": [18, 23, 30, 39, 51, 65, 81]
        },
        
        "description": "Low pressure drop fill for applications with air-side limitations or dirty water"
    }
}

# ============================================================================
# TOWER TYPE DATABASE
# ============================================================================

TOWER_TYPES = {
    "crossflow": {
        "name": "Crossflow",
        "air_water_contact": "Perpendicular",
        "typical_pressure_drop_factor": 1.0,
        "air_distribution": "Side inlet",
        "fill_utilization": 0.85,
        "description": "Air flows horizontally, water flows vertically downward"
    },
    "counterflow_induced": {
        "name": "Counterflow (Induced Draft)",
        "air_water_contact": "Parallel counter-current",
        "typical_pressure_drop_factor": 1.3,  # Higher pressure drop
        "air_distribution": "Bottom inlet, top fan",
        "fill_utilization": 0.95,
        "description": "Air flows upward against downward water flow, fan on top"
    },
    "counterflow_forced": {
        "name": "Counterflow (Forced Draft)",
        "air_water_contact": "Parallel counter-current",
        "typical_pressure_drop_factor": 1.2,
        "air_distribution": "Bottom fan, top outlet",
        "fill_utilization": 0.92,
        "description": "Air flows upward, fan at bottom"
    }
}


# ============================================================================
# REPORTING / GEOMETRY HELPERS
# ============================================================================

def calculate_open_air_area(face_area_m2, free_area_fraction):
    """Net open area through fill based on an assumed free-area fraction."""
    try:
        return max(float(face_area_m2), 0.0) * max(min(float(free_area_fraction), 1.0), 0.0)
    except Exception:
        return None


def pressure_drop_method_label(fill_data, tower_type, use_pdf_data):
    """Readable explanation of the pressure-drop basis used in the calculation."""
    name = str(fill_data.get("name", ""))
    if name == "Brentwood CT1200AT (New from PDF)" and use_pdf_data and str(tower_type).startswith("counterflow"):
        return "Brentwood published CT1200AT/CF1200-type equation; fill ΔP used directly, then non-fill losses are added"
    if name.startswith("Brentwood CF1900SB/MA") and use_pdf_data and str(tower_type).startswith("counterflow"):
        return "Brentwood published CF1900SB/MA equation; fill ΔP used directly, then non-fill losses are added"
    if name.startswith("Brentwood XF75") and use_pdf_data and tower_type == "crossflow":
        return "Brentwood XF75 SI equation; fill ΔP used directly, then non-fill losses are added"
    return "Legacy/default curve basis; tower pressure-drop factor is applied as an empirical allowance, then non-fill losses are added"


# ============================================================================
# BRENTWOOD XF75 CROSSFLOW SI DATA (2018)
# ============================================================================

# SI charts in XF75 Performance Data_2018.pdf:
# - Thermal performance pages: 10-13 (KaH/L vs L/G for fixed fill height and AT)
# - Pressure drop pages: 14-18 (Pa vs air velocity for fixed AT and water loading)
#
# The app keeps the user input as L/G.  For the XF75 thermal correlation/table
# we build an interpolation grid from the published SI curve family form:
#   KaH/L = 1.706 * (G/L)^0.822 * H^0.178 * AT^0.822
# where H and AT are in meters.  Since UI uses L/G, G/L = 1/(L/G).
#
# The selected/published values below match the SI graph families.

XF75_SI_FILL_HEIGHT_OPTIONS_M = [1.829, 2.438, 3.657, 4.877]
XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M = [0.610, 0.914, 1.219, 1.524, 1.829]

# XF75 catalogue/product block details (not thermal curve data).
# Catalogue dimensions: D = air travel depth, W = module/stacking width, L = vertical fill height.
XF75_CATALOG_SURFACE_AREA_M2_M3 = 167.4
XF75_CATALOG_SHEET_SPACING_MM = 19.0
XF75_CATALOG_SHEETS_PER_FOOT = 16
XF75_CATALOG_MIN_D_M = 0.305
XF75_CATALOG_MIN_W_M = 0.153
XF75_CATALOG_MIN_L_M = 0.915
XF75_CATALOG_MAX_D_M = 0.610
XF75_CATALOG_MAX_W_M = 0.305
XF75_CATALOG_MAX_L_M = 3.658
XF75_CATALOG_STANDARD_D_M = 0.610
XF75_CATALOG_STANDARD_W_M = 0.305
XF75_CATALOG_STANDARD_L_OPTIONS_M = [1.829, 2.439, 3.048, 3.658]

XF75_SI_LG_GRID = [0.10, 0.15, 0.20, 0.25, 0.30, 0.40, 0.50, 0.60, 0.70, 0.80, 0.90,
                   1.00, 1.50, 2.00, 2.50, 3.00]
XF75_SI_WATER_LOADING_RANGE = (20.0, 70.0)       # m³/(h·m²), pressure-drop graph range
XF75_SI_AIR_VELOCITY_RANGE = (1.0, 4.0)          # m/s, pressure-drop graph range
XF75_SI_LG_RANGE = (0.10, 3.00)                  # from thermal graph axis


# Published XF75 SI chart family combinations from Brentwood pages 10-13.
# These are not arbitrary: each height chart has only certain AT curves printed.
XF75_SI_PUBLISHED_HEIGHT_AT_COMBOS = {
    1.829: [0.610, 0.914, 1.219],
    2.438: [0.914, 1.219, 1.524],
    3.657: [1.219, 1.524, 1.829],
    4.877: [1.219, 1.524, 1.829],
}


def _range_status(value, lower, upper, units=""):
    """Return (ok_bool, readable_status) for range checks."""
    ok = lower <= value <= upper
    status = "OK" if ok else "OUTSIDE"
    return ok, f"{status}: {value:.3g}{units} vs Brentwood graph range {lower:g}–{upper:g}{units}"


def xf75_graph_validity_checks(l_over_g, water_loading, air_face_velocity, fill_height_m, air_travel_depth_m):
    """Return Brentwood XF75 SI graph-range checks for UI and reports.

    Ranges are taken from the SI Brentwood XF75 performance pages:
    - Thermal: L/G axis about 0.10–3.00; H/AT chart families.
    - Pressure drop: water loading 20–70 m³/(h·m²), air velocity 1–4 m/s, AT 0.610–1.829 m.
    """
    checks = []
    warnings = []

    def add(label, ok, message):
        line = f"{label}: {message}"
        checks.append((label, ok, message))
        if not ok:
            warnings.append(line)

    ok, msg = _range_status(float(l_over_g), *XF75_SI_LG_RANGE, "")
    add("L/G", ok, msg)

    ok, msg = _range_status(float(water_loading), *XF75_SI_WATER_LOADING_RANGE, " m³/h·m²")
    add("Water loading", ok, msg)

    ok, msg = _range_status(float(air_face_velocity), *XF75_SI_AIR_VELOCITY_RANGE, " m/s")
    add("Air face velocity", ok, msg)

    h = float(fill_height_m)
    at = float(air_travel_depth_m)
    h_min, h_max = min(XF75_SI_FILL_HEIGHT_OPTIONS_M), max(XF75_SI_FILL_HEIGHT_OPTIONS_M)
    at_min, at_max = min(XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M), max(XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M)

    ok, msg = _range_status(h, h_min, h_max, " m")
    add("Fill height H", ok, msg)

    ok, msg = _range_status(at, at_min, at_max, " m")
    add("Air travel depth AT", ok, msg)

    # Check whether the selected H/AT is exactly one of the printed SI thermal-chart curves.
    nearest_h = min(XF75_SI_FILL_HEIGHT_OPTIONS_M, key=lambda x: abs(x - h))
    nearest_at = min(XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M, key=lambda x: abs(x - at))
    exact_h = abs(h - nearest_h) < 0.002
    exact_at = abs(at - nearest_at) < 0.002
    combo_ok = exact_h and exact_at and any(abs(nearest_at - a) < 0.002 for a in XF75_SI_PUBLISHED_HEIGHT_AT_COMBOS.get(nearest_h, []))
    combo_msg = (
        f"OK: H={h:.3f} m and AT={at:.3f} m are a published XF75 SI thermal-chart combination"
        if combo_ok else
        f"INTERPOLATED/EXTRAPOLATED: H={h:.3f} m and AT={at:.3f} m are not an exact printed Brentwood SI thermal-chart combination"
    )
    add("Published H/AT combination", combo_ok, combo_msg)

    return {
        "all_ok": len(warnings) == 0,
        "checks": checks,
        "warnings": warnings,
    }


def xf75_kah_over_l_si(l_over_g: float, fill_height_m: float, air_travel_depth_m: float) -> float:
    """Return Brentwood XF75 crossflow KaH/L using SI curve-family equation.

    This is the thermal NTU-like term for the selected fill height H and air travel depth AT.
    The user-facing input remains L/G; internally the SI relation uses G/L.
    """
    lg = max(float(l_over_g), 1e-9)
    h = max(float(fill_height_m), 1e-9)
    at = max(float(air_travel_depth_m), 1e-9)
    g_over_l = 1.0 / lg
    return 1.706 * (g_over_l ** 0.822) * (h ** 0.178) * (at ** 0.822)


def xf75_kah_over_l_from_table(l_over_g: float, fill_height_m: float, air_travel_depth_m: float) -> float:
    """2D/3D style interpolation for XF75 KaH/L.

    Brentwood publishes discrete H and AT chart families.  We generate values at those
    exact chart families and interpolate in log-space across H, AT and L/G.  This behaves
    like a digitized curve table while preserving smoothness between published families.
    """
    import numpy as _np

    lg = float(l_over_g)
    h = float(fill_height_m)
    at = float(air_travel_depth_m)

    # Clamp for interpolation; separate warnings are generated elsewhere.
    lg_c = min(max(lg, min(XF75_SI_LG_GRID)), max(XF75_SI_LG_GRID))
    h_c = min(max(h, min(XF75_SI_FILL_HEIGHT_OPTIONS_M)), max(XF75_SI_FILL_HEIGHT_OPTIONS_M))
    at_c = min(max(at, min(XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M)), max(XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M))

    # Because published families follow a power law, log-space interpolation is more stable.
    x_grid = _np.log(_np.array(XF75_SI_LG_GRID, dtype=float))
    x = _np.log(lg_c)

    def interp_lg_for(hv, atv):
        vals = [xf75_kah_over_l_si(v, hv, atv) for v in XF75_SI_LG_GRID]
        return float(_np.exp(_np.interp(x, x_grid, _np.log(vals))))

    # Interpolate over AT for each H
    at_grid = _np.log(_np.array(XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M, dtype=float))
    at_x = _np.log(at_c)

    values_by_h = []
    for hv in XF75_SI_FILL_HEIGHT_OPTIONS_M:
        vals_at = [interp_lg_for(hv, atv) for atv in XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M]
        values_by_h.append(float(_np.exp(_np.interp(at_x, at_grid, _np.log(vals_at)))))

    h_grid = _np.log(_np.array(XF75_SI_FILL_HEIGHT_OPTIONS_M, dtype=float))
    h_x = _np.log(h_c)
    return float(_np.exp(_np.interp(h_x, h_grid, _np.log(values_by_h))))


def xf75_crossflow_geometry_from_session(default_height=1.829, default_at=0.914, default_width=6.0, default_banks=2):
    """Read crossflow geometry from Streamlit session state safely.

    bank_count = 1 for a single-sided crossflow fill bank, or 2 for the common
    induced-draft crossflow arrangement with fill banks on both sides of a central plenum/fan.
    """
    h = float(st.session_state.get('crossflow_fill_height_m', default_height))
    at = float(st.session_state.get('crossflow_air_travel_depth_m', default_at))
    width = float(st.session_state.get('crossflow_stack_width_m', default_width))
    bank_count = int(st.session_state.get('crossflow_bank_count', default_banks))
    if bank_count not in (1, 2):
        bank_count = default_banks
    return h, at, width, bank_count


def xf75_catalog_module_estimate(fill_height_m: float, air_travel_depth_m: float, stack_width_m: float):
    """Estimate how the selected XF75 geometry maps to catalogue media pack dimensions.

    Catalogue D = air travel depth, W = module width/stacking direction, L = vertical fill height.
    Standard catalogue pack: D=0.610 m, W=0.305 m, L=1.829/2.439/3.048/3.658 m.
    Thermal charts also include 4.877 m (16 ft), which would normally be achieved by vertical stacking or a custom arrangement.
    This is a practical packing estimate, not a structural bill of materials.
    """
    import math as _math
    h = max(float(fill_height_m or 0), 0.0)
    at = max(float(air_travel_depth_m or 0), 0.0)
    w = max(float(stack_width_m or 0), 0.0)
    depth_modules = _math.ceil(at / XF75_CATALOG_STANDARD_D_M) if at > 0 else 0
    width_modules = _math.ceil(w / XF75_CATALOG_STANDARD_W_M) if w > 0 else 0
    height_near_standard = min(XF75_CATALOG_STANDARD_L_OPTIONS_M, key=lambda x: abs(x - h)) if h > 0 else None
    vertical_layers = _math.ceil(h / XF75_CATALOG_MAX_L_M) if h > 0 else 0
    return {
        "catalog_D_air_travel_depth_m": XF75_CATALOG_STANDARD_D_M,
        "catalog_W_module_width_m": XF75_CATALOG_STANDARD_W_M,
        "catalog_L_standard_heights_m": XF75_CATALOG_STANDARD_L_OPTIONS_M,
        "depth_modules_est": depth_modules,
        "width_modules_est": width_modules,
        "vertical_layers_est": vertical_layers,
        "nearest_catalog_L_m": height_near_standard,
    }

# ============================================================================
# PSYCHROMETRIC FUNCTIONS WITH DRY BULB SUPPORT
# ============================================================================

def saturation_pressure(temp_C):
    """Calculate saturation vapor pressure in kPa"""
    T = temp_C + 273.15
    if temp_C >= 0:
        return 0.61121 * np.exp((18.678 - temp_C/234.5) * (temp_C/(257.14 + temp_C)))
    else:
        return 0.61115 * np.exp((23.036 - temp_C/333.7) * (temp_C/(279.82 + temp_C)))

def humidity_ratio_from_wb(db, wb, pressure=101.325):
    """Calculate humidity ratio from dry bulb and wet bulb temperatures"""
    Pws_wb = saturation_pressure(wb)
    Ws_wb = 0.62198 * Pws_wb / (pressure - Pws_wb)
    
    h_fg = 2501.0
    Cp_air = 1.006
    Cp_vapor = 1.86
    
    W = ((h_fg - Cp_vapor * wb) * Ws_wb - Cp_air * (db - wb)) / (h_fg + Cp_vapor * db - 4.186 * wb)
    return max(W, 0.0001)

def relative_humidity_from_wb(db, wb, pressure=101.325):
    """Calculate relative humidity from dry bulb and wet bulb temperatures"""
    W = humidity_ratio_from_wb(db, wb, pressure)
    Pws_db = saturation_pressure(db)
    Ws_db = 0.62198 * Pws_db / (pressure - Pws_db)
    return (W / Ws_db) * 100

def enthalpy_air(db, W):
    """Calculate enthalpy of moist air in kJ/kg dry air"""
    Cp_air = 1.006
    Cp_vapor = 1.86
    h_fg = 2501.0
    return Cp_air * db + W * (h_fg + Cp_vapor * db)

def air_density_calc(db, wb, altitude=0):
    """Calculate air density considering altitude and humidity"""
    # Atmospheric pressure at altitude
    P_atm = 101.325 * (1 - 0.0000225577 * altitude) ** 5.25588  # kPa
    
    # Humidity ratio
    W = humidity_ratio_from_wb(db, wb, P_atm)
    
    # Gas constant for dry air
    R_da = 0.28705  # kJ/kg·K
    
    # Temperature in Kelvin
    T_K = db + 273.15
    
    # Density using ideal gas law with humidity correction
    rho = (P_atm * 1000) / (R_da * 1000 * T_K * (1 + 1.609 * W))
    
    return rho

# ============================================================================
# ENHANCED CALCULATION FUNCTIONS WITH TOWER TYPE SUPPORT
# ============================================================================

def calculate_pressure_drop_with_tower_type(fill_data, tower_type, air_face_velocity, water_loading, fill_depth, use_pdf_data=False):
    """Calculate pressure drop considering tower type"""
    L_prime = water_loading / 3.6  # Convert to kg/(s·m²)
    
    # Get base pressure drop from fill curve
    if fill_data["name"] == "Brentwood CT1200AT (New from PDF)" and use_pdf_data and tower_type.startswith("counterflow"):
        # Use Brentwood's published CT1200AT pressure drop equation (from the performance sheet).
        #
        # IP correlation form:
        #   DP = [4.6192e-6 * vel^(1.7443) + qa*(4.9355e-9 * vel^(2.3711))] * (0.1513 + 0.2852*ht)
        # where:
        #   vel = air velocity (ft/min)
        #   qa  = water loading (gpm/ft²)
        #   ht  = fill height (ft)
        #   DP  = inches of water gauge (in.wg)
        #
        # Convert SI inputs to the IP units expected by the correlation.
        vel_fpm = air_face_velocity * 196.850394  # m/s -> ft/min
        qa_gpm_ft2 = water_loading * (4.402867 / 10.763910416709722)  # (m³/h·m²) -> (gpm/ft²)
        ht_ft = fill_depth / 0.3048

        dp_inwg = (
            (4.6192e-6 * (vel_fpm ** 1.7443)) +
            (qa_gpm_ft2 * (4.9355e-9 * (vel_fpm ** 2.3711)))
        ) * (0.1513 + 0.2852 * ht_ft)

        # Convert in.wg to Pa
        delta_P_base = dp_inwg * 249.0889  # Pa

        # No additional velocity² or height scaling needed (already included in the correlation)
        velocity_factor = 1.0

    elif fill_data["name"] == "Brentwood CF1900SB/MA (2017 Test Cell)":
        # Use Brentwood's published CF1900SB/MA pressure drop equation (from the 2017 sheet).
        #
        # IP correlation form:
        #   DP = [2.2470e-6 * vel^(1.7897) + qa*(4.8967e-8 * vel^(1.9362))] * (0.1984 + 0.3281*ht)
        # where:
        #   vel = air velocity (ft/min)
        #   qa  = water loading (gpm/ft²)
        #   ht  = fill height (ft)
        #   DP  = inches of water gauge (in.wg)
        vel_fpm = air_face_velocity * 196.850394  # m/s -> ft/min
        qa_gpm_ft2 = water_loading * (4.402867 / 10.763910416709722)  # (m³/h·m²) -> (gpm/ft²)
        ht_ft = fill_depth / 0.3048

        dp_inwg = (
            (2.2470e-6 * (vel_fpm ** 1.7897)) +
            (qa_gpm_ft2 * (4.8967e-8 * (vel_fpm ** 1.9362)))
        ) * (0.1984 + 0.3281 * ht_ft)

        delta_P_base = dp_inwg * 249.0889  # in.wg -> Pa
        velocity_factor = 1.0

    elif str(fill_data.get("name","")).startswith("Brentwood XF75") and use_pdf_data and tower_type == "crossflow":
        # Brentwood XF75 crossflow pressure drop, SI form from XF75 Performance Data_2018.pdf.
        #
        #   ΔP (Pa) = [12.879 * V^1.7532 + QA*(6.2001e-3 * V^2) + 0.32229*V] * AT
        # where:
        #   V  = air face velocity through fill (m/s)
        #   QA = water loading over plan area (m³/(h·m²))
        #   AT = horizontal air travel depth through fill (m)
        V = air_face_velocity
        QA = water_loading
        _, AT, _, _ = xf75_crossflow_geometry_from_session(default_height=fill_depth, default_at=0.914, default_width=6.0)
        if AT <= 0:
            AT = 0.914

        delta_P_base = (12.879 * (V ** 1.7532) + QA * (6.2001e-3 * (V ** 2)) + 0.32229 * V) * AT
        velocity_factor = 1.0

    elif fill_data["name"] == "Brentwood CT1200AT (New from PDF)" and use_pdf_data:

        # Fallback: use digitized PDF table data (works for non-counterflow cases too)
        if fill_depth <= 0.61:  # 24" or less
            delta_P_base = np.interp(
                L_prime * 3.6,
                [x * 8 for x in fill_data["performance_data"]["L_G"]],
                fill_data["performance_data"]["delta_P_base_24"]
            )
        elif fill_depth <= 0.915:  # 36" or less
            delta_P_base = np.interp(
                L_prime * 3.6,
                [x * 8 for x in fill_data["performance_data"]["L_G"]],
                fill_data["performance_data"]["delta_P_base_36"]
            )
        else:  # 48" or more
            delta_P_base = np.interp(
                L_prime * 3.6,
                [x * 8 for x in fill_data["performance_data"]["L_G"]],
                fill_data["performance_data"]["delta_P_base_48"]
            )
        velocity_factor = (air_face_velocity / 2.5) ** 2
    else:
        # Use standard data for other fills
        delta_P_base = np.interp(
            L_prime * 3.6,  # Temporary conversion for interpolation
            [x * 8 for x in fill_data["performance_data"]["L_G"]],  # Approximate scaling
            fill_data["performance_data"]["delta_P_base"]
        )
    
    # Adjust for actual face velocity (ΔP ∝ velocity²) when using the legacy curve method
    if 'velocity_factor' not in locals():
        velocity_factor = (air_face_velocity / 2.5) ** 2
    
    # Tower pressure-drop factor is retained only for legacy/default curves.
    # Published Brentwood equations already calculate fill pressure drop for the selected geometry.
    tower_factor = TOWER_TYPES[tower_type]["typical_pressure_drop_factor"]

    # Fill pressure drop
    # For equation-based correlations that already include full geometry, do NOT multiply again by fill height
    # and do NOT multiply by the generic 1.2/1.3 tower factor.
    # - CT1200AT/CF1200-type counterflow ΔP equation includes fill height ht.
    # - CF1900SB/MA ΔP equation includes fill height ht.
    # - XF75 crossflow ΔP equation includes air travel depth AT.
    # Legacy digitized curves are kept as before and are scaled by fill_depth and tower_factor.
    is_ct1200_equation = (fill_data.get('name') == 'Brentwood CT1200AT (New from PDF)' and use_pdf_data and tower_type.startswith('counterflow'))
    is_cf1900_equation = (str(fill_data.get('name','')).startswith('Brentwood CF1900SB/MA') and use_pdf_data and str(tower_type).startswith('counterflow'))
    is_xf75_crossflow_equation = (str(fill_data.get('name','')).startswith('Brentwood XF75') and use_pdf_data and tower_type == 'crossflow')
    is_published_dp_equation = is_ct1200_equation or is_cf1900_equation or is_xf75_crossflow_equation
    geometry_multiplier = 1.0 if is_published_dp_equation else fill_depth
    applied_tower_factor = 1.0 if is_published_dp_equation else tower_factor
    fill_pressure_drop = delta_P_base * velocity_factor * geometry_multiplier * applied_tower_factor

    # Additional losses based on tower type
    if tower_type == "counterflow_induced":
        # Inlet, eliminators, fan inlet losses (matching supplier's SAA15)
        additional_losses = {
            "inlet_louver": 15,  # Pa (K=3.0)
            "eliminators": 10,   # Pa (K=2.0)
            "fan_inlet": 2,      # Pa (K=0.3)
            "stack_exit": 2,
            "rain_zone": 5
        }
    else:  # crossflow
        additional_losses = {
            "inlet_louver": 10,
            "eliminators": 8,
            "fan_inlet": 1,
            "stack_exit": 1,
            "rain_zone": 3
        }
    
    total_static_pressure = fill_pressure_drop + sum(additional_losses.values())
    
    return {
        "fill_pressure_drop": fill_pressure_drop,
        "total_static_pressure": total_static_pressure,
        "additional_losses": additional_losses,
        "other_losses_total": sum(additional_losses.values()),
        "tower_factor": applied_tower_factor,
        "legacy_tower_factor": tower_factor,
        "uses_published_dp_equation": is_published_dp_equation,
        "pressure_drop_method": pressure_drop_method_label(fill_data, tower_type, use_pdf_data),
        "delta_P_base": delta_P_base
    }

def calculate_KaL_with_tower_type(fill_data, L_over_G, tower_type, fill_depth=None, use_pdf_data=False):
    """Calculate Ka/L considering tower type adjustments.

    Notes:
    - For CT1200AT + PDF mode + counterflow tower types, Brentwood's published KaV/L equation is used
      (with internal unit conversions). The solver downstream still expects Ka/L, so we convert:
        Ka/L = (KaV/L) / H
      where H is fill height in meters.
    - For all other cases, the existing interpolated Ka/L curves are used (unchanged).
    """
    use_brentwood_equation = (
        fill_data.get("name") == "Brentwood CT1200AT (New from PDF)"
        and use_pdf_data
        and isinstance(tower_type, str)
        and tower_type.startswith("counterflow")
        and fill_depth is not None
        and fill_depth > 0
    )

    if use_brentwood_equation:
        # Brentwood CT1200AT KaV/L correlation (IP form on the performance sheet):
        #   KaV/L = 0.967 * (L/G)^(-0.779) * H^(0.632) , with H in nominal feet.
        H_ft = fill_depth / 0.3048
        KaV_over_L = 0.967 * (L_over_G ** (-0.779)) * (H_ft ** 0.632)

        # Convert to Ka/L per meter for the existing NTU calculation: NTU = (Ka/L) * H
        Ka_over_L = KaV_over_L / fill_depth
        return Ka_over_L



    use_cf1900_equation = (
        str(fill_data.get("name","")).startswith("Brentwood CF1900SB/MA")
        and use_pdf_data
        and isinstance(tower_type, str)
        and fill_depth is not None
        and fill_depth > 0
    )

    if use_cf1900_equation:
        # Brentwood CF1900SB/MA KaV/L correlation (from the performance sheet image):
        #   KaV/L = 0.696 * (L/G)^(-0.707) * H^(0.714) , with H in feet.
        # KaV/L is dimensionless; convert to Ka/L per meter to preserve NTU = (Ka/L)*H:
        #   Ka/L = (KaV/L) / H_m
        H_ft = fill_depth / 0.3048
        KaV_over_L = 0.696 * (L_over_G ** (-0.707)) * (H_ft ** 0.714)
        Ka_over_L = KaV_over_L / fill_depth
        return Ka_over_L

    use_xf75_equation = (
        str(fill_data.get("name","")).startswith("Brentwood XF75")
        and use_pdf_data
        and tower_type == "crossflow"
        and fill_depth is not None
        and fill_depth > 0
    )

    if use_xf75_equation:
        # Brentwood XF75 crossflow thermal performance, SI curve table.
        # The XF75 SI graphs are families of KaH/L vs L/G for fixed fill height H and
        # air travel depth AT. The app keeps the UI in L/G and uses a log-space
        # interpolation table built from the Brentwood SI curve family.
        H_m, AT_m, _, _ = xf75_crossflow_geometry_from_session(default_height=fill_depth, default_at=0.914, default_width=6.0)
        if H_m <= 0:
            H_m = fill_depth
        if AT_m <= 0:
            AT_m = 0.914

        KaH_over_L = xf75_kah_over_l_from_table(L_over_G, H_m, AT_m)
        # Existing solver expects Ka/L and then forms NTU = (Ka/L)*H.
        # Therefore Ka/L = KaH/L / H.
        Ka_over_L = KaH_over_L / H_m
        return Ka_over_L

    # -------------------------------------------------------------------------
    # Default (existing) behavior: interpolate Ka/L from fill curve
    # -------------------------------------------------------------------------
    Ka_over_L = np.interp(
        L_over_G,
        fill_data["performance_data"]["L_G"],
        fill_data["performance_data"]["Ka_L"]
    )

    # Adjust for tower type (counterflow typically has better utilization)
    if tower_type.startswith("counterflow"):
        # Counterflow has better contact efficiency
        efficiency_factor = TOWER_TYPES[tower_type]["fill_utilization"] / 0.85
        Ka_over_L *= efficiency_factor

    return Ka_over_L

# ============================================================================
# MAIN CALCULATION FUNCTION WITH DRY BULB & ALTITUDE SUPPORT

# ============================================================================

def solve_cooling_tower_enhanced(L, G, T_hot, T_cold_target, Twb, Tdb, fill_type, 
                                 tower_type, fill_depth, face_area, altitude=0, use_pdf_data=False):
    """
    Enhanced cooling tower solver with dry bulb temperature and altitude support
    """
    fill_data = BRENTWOOD_FILLS[fill_type]
    tower_data = TOWER_TYPES[tower_type]
    
    # Get adjusted Ka/L
    L_over_G = L / G

    # Crossflow geometry is different from counterflow geometry:
    #   air face area  = fill height × stacked fill width
    #   water plan area = air travel depth × stacked fill width
    #   fill volume    = fill height × air travel depth × stacked fill width
    # For counterflow, the existing face_area × fill_depth method is preserved.
    crossflow_geometry = {}
    if tower_type == "crossflow":
        cf_height, cf_at, cf_width, cf_bank_count = xf75_crossflow_geometry_from_session(default_height=fill_depth, default_at=0.914, default_width=6.0)
        cf_height = max(cf_height, 1e-6)
        cf_at = max(cf_at, 1e-6)
        cf_width = max(cf_width, 1e-6)
        cf_bank_count = 2 if int(cf_bank_count) == 2 else 1
        fill_depth = cf_height

        # For a dual-bank induced-draft crossflow tower, the fan draws air through
        # two identical fill banks on opposite sides of the central plenum. Total
        # air face area, water plan area and fill volume are multiplied by bank_count.
        air_face_area_single_bank = cf_height * cf_width
        water_plan_area_single_bank = cf_at * cf_width
        fill_volume_single_bank = cf_height * cf_at * cf_width
        air_face_area = cf_bank_count * air_face_area_single_bank
        water_plan_area = cf_bank_count * water_plan_area_single_bank
        fill_volume_calc = cf_bank_count * fill_volume_single_bank
        crossflow_geometry = {
            "crossflow_fill_height_m": cf_height,
            "crossflow_air_travel_depth_m": cf_at,
            "crossflow_stack_width_m": cf_width,
            "crossflow_bank_count": cf_bank_count,
            "air_face_area_single_bank_m2": air_face_area_single_bank,
            "water_plan_area_single_bank_m2": water_plan_area_single_bank,
            "fill_volume_single_bank_m3": fill_volume_single_bank,
            "air_face_area_m2": air_face_area,
            "water_plan_area_m2": water_plan_area,
        }
    else:
        air_face_area = face_area
        water_plan_area = face_area
        fill_volume_calc = face_area * fill_depth

    Ka_over_L = calculate_KaL_with_tower_type(fill_data, L_over_G, tower_type, fill_depth, use_pdf_data)
    
    # Total heat transfer coefficient (K value)
    Ka = Ka_over_L * L  # kW/°C
    
    # Air properties with dry bulb and altitude
    air_density = air_density_calc(Tdb, Twb, altitude)  # kg/m³
    air_flow_volumetric = G / air_density  # m³/s
    air_face_velocity = air_flow_volumetric / air_face_area  # m/s
    
    # Water loading over the wetted plan area
    water_loading = (L * 3.6) / water_plan_area  # m³/h·m²
    
    # Calculate pressure drop with tower type consideration
    pressure_results = calculate_pressure_drop_with_tower_type(
        fill_data, tower_type, air_face_velocity, water_loading, fill_depth, use_pdf_data
    )
    
    # Calculate hydraulic properties
    water_velocity_ms = (water_loading / 3600.0) / fill_data["water_passage_area"]
    water_viscosity = 1e-6
    film_reynolds = (water_velocity_ms * fill_data["water_film_thickness"] * 1e-3) / water_viscosity
    air_viscosity = 1.5e-5
    air_reynolds = (air_face_velocity * fill_data["hydraulic_diameter"] * 1e-3) / air_viscosity
    
    # Assess fouling risk
    risk_score = 0
    if fill_data["hydraulic_diameter"] < 10:
        risk_score += 2
    elif fill_data["hydraulic_diameter"] < 12:
        risk_score += 1
    if water_velocity_ms < 0.05:
        risk_score += 2
    elif water_velocity_ms < 0.1:
        risk_score += 1
    
    risk_level = "Low" if risk_score < 2 else "Moderate" if risk_score < 4 else "High"
    
    # Merkel number (NTU) - adjusted for tower type
    NTU = Ka_over_L * fill_depth
    if tower_type.startswith("counterflow"):
        # Counterflow typically achieves 5-10% better NTU utilization
        NTU *= 1.05
    
    # Achieved cold water temperature (simplified Merkel solution)
    # For counterflow, use more accurate approach
    if tower_type.startswith("counterflow"):
        # Counterflow typically has better approach for same NTU
        approach_factor = 0.95
    else:
        approach_factor = 1.0
    
    T_cold_achieved = Twb + (T_hot - Twb) * np.exp(-NTU * approach_factor)
    
    # Ensure realistic temperature
    T_cold_achieved = max(T_cold_achieved, Twb + 0.5)
    T_cold_achieved = min(T_cold_achieved, T_hot - 0.5)
    
    Q_achieved = L * 4.186 * (T_hot - T_cold_achieved)
    
    # Fill volume and surface area
    fill_volume = fill_volume_calc
    total_surface_area = fill_volume * fill_data["surface_area"]
    
    # Calculate fan power
    fan_efficiency = 0.78  # 78% as per supplier SAA15
    transmission_efficiency = 1.0
    fan_power = (air_flow_volumetric * pressure_results["total_static_pressure"]) / \
                (fan_efficiency * transmission_efficiency * 1000)  # kW
    
    # Calculate relative humidity
    P_atm = 101.325 * (1 - 0.0000225577 * altitude) ** 5.25588
    RH = relative_humidity_from_wb(Tdb, Twb, P_atm)
    
    # Geometry input metadata for report clarity
    geometry_input = dict(st.session_state.get("geometry_input", {})) if hasattr(st, "session_state") else {}

    # Operating warnings
    operating_warnings = []
    xf75_graph_checks = None

    # For Brentwood XF75 crossflow PDF mode, use the actual SI graph envelopes rather than generic fill limits.
    # This is what protects the user from relying on extrapolated Brentwood data.
    if tower_type == "crossflow" and str(fill_data.get("name", "")).startswith("Brentwood XF75") and use_pdf_data:
        cf_height, cf_at, cf_width, cf_bank_count = xf75_crossflow_geometry_from_session(default_height=fill_depth, default_at=0.914, default_width=6.0)
        xf75_graph_checks = xf75_graph_validity_checks(L_over_G, water_loading, air_face_velocity, cf_height, cf_at)
        for w in xf75_graph_checks["warnings"]:
            operating_warnings.append("XF75 Brentwood graph range: " + w)
    else:
        if water_loading > fill_data["max_water_loading"]:
            operating_warnings.append(f"Water loading exceeds maximum ({fill_data['max_water_loading']} m³/h·m²)")
        if water_loading < fill_data["min_water_loading"]:
            operating_warnings.append(f"Water loading below minimum ({fill_data['min_water_loading']} m³/h·m²)")
        if air_face_velocity > fill_data["max_air_velocity"]:
            operating_warnings.append(f"Air face velocity exceeds maximum ({fill_data['max_air_velocity']} m/s)")
    
    # Tower type specific warnings
    if tower_type == "counterflow_induced" and air_face_velocity > 2.5:
        operating_warnings.append("High air velocity for induced draft - ensure proper plenum design")
    
    return {
        # Basic identification
        "fill_type": fill_type,
        "fill_name": fill_data["name"],
        "tower_type": tower_type,
        "tower_name": tower_data["name"],
        "use_pdf_data": use_pdf_data,
        
        # Temperatures and heat transfer
        "T_hot": T_hot,
        "T_cold_achieved": T_cold_achieved,
        "T_cold_target": T_cold_target,
        "Twb": Twb,
        "Tdb": Tdb,
        "RH": RH,
        "Q_achieved": Q_achieved,
        "Q_target": L * 4.186 * (T_hot - T_cold_target),
        "approach": T_cold_achieved - Twb,
        "cooling_range": T_hot - T_cold_achieved,
        
        # Flow parameters
        "L": L,
        "G": G,
        "L_over_G": L_over_G,
        "water_loading": water_loading,
        "air_density": air_density,
        "air_flow_volumetric": air_flow_volumetric,
        "air_face_velocity": air_face_velocity,
        
        # Geometry and sizing
        "fill_depth": fill_depth,
        "face_area": air_face_area,
        "water_plan_area": water_plan_area,
        "input_face_area": face_area,
        "fill_volume": fill_volume,
        "total_surface_area": total_surface_area,
        "crossflow_geometry": crossflow_geometry,
        
        # Hydraulic properties
        "water_velocity": water_velocity_ms,
        "film_reynolds": film_reynolds,
        "air_reynolds": air_reynolds,
        "water_film_thickness": fill_data["water_film_thickness"],
        
        # Performance parameters - INCLUDING K VALUE
        "NTU": NTU,
        "Ka_over_L": Ka_over_L,
        "KaH_over_L": Ka_over_L * fill_depth,
        "KaV_over_L": Ka_over_L * fill_depth,
        "Ka": Ka,  # This is the K value (heat transfer coefficient in kW/°C)
        "K_value": Ka,  # Added for clarity
        
        # Pressure drop and fan
        "fill_pressure_drop": pressure_results["fill_pressure_drop"],
        "total_static_pressure": pressure_results["total_static_pressure"],
        "additional_losses": pressure_results["additional_losses"],
        "other_losses_total": pressure_results.get("other_losses_total", sum(pressure_results["additional_losses"].values())),
        "fan_power": fan_power,
        "delta_P_base": pressure_results.get("delta_P_base", 0),
        "pressure_drop_method": pressure_results.get("pressure_drop_method", ""),
        "uses_published_dp_equation": pressure_results.get("uses_published_dp_equation", False),
        "tower_pressure_factor_applied": pressure_results.get("tower_factor", 1.0),
        "legacy_tower_pressure_factor": pressure_results.get("legacy_tower_factor", TOWER_TYPES[tower_type]["typical_pressure_drop_factor"]),
        
        # Assessments
        "fouling_risk": {"risk_score": risk_score, "risk_level": risk_level},
        "operating_warnings": operating_warnings,
        "xf75_graph_checks": xf75_graph_checks,
        
        # Fill characteristics
        "surface_area_density": fill_data["surface_area"],
        "hydraulic_diameter": fill_data["hydraulic_diameter"],
        "flute_angle": fill_data["flute_angle"],
        "free_area_fraction": fill_data["free_area_fraction"],
        
        # Tower characteristics
        "tower_efficiency_factor": tower_data["fill_utilization"],
        
        # Atmospheric conditions
        "altitude": altitude,
        "air_density_calc": air_density
    }

# ============================================================================
# SUPPLIER SAA15 DESIGN VALIDATION FUNCTION
# ============================================================================

def validate_with_saa15_supplier_design():
    """
    Run validation against supplier's SAA15 design with CF1200
    Returns the results for comparison
    """
    # Supplier's SAA15 parameters from image
    supplier_inputs = {
        "L": 114,  # kg/s (actual water flow)
        "G": 49.28,  # kg/s (calculated from L/G=2.313)
        "T_hot": 40.0,  # °C (from range 5°C and T_cold=35°C)
        "T_cold_target": 35.0,  # °C
        "Twb": 30.0,  # °C (assumed from your input)
        "Tdb": 33.0,  # °C (estimated for 60% RH)
        "fill_type": "CF1200",
        "tower_type": "counterflow_induced",
        "fill_depth": 0.75,  # m
        "face_area": 12.96,  # m² (3.6m x 3.6m)
        "altitude": 0
    }
    
    # Run calculation
    results = solve_cooling_tower_enhanced(**supplier_inputs)
    
    # Supplier's claimed results from image
    supplier_claimed = {
        "T_cold_achieved": 35.0,  # °C
        "exit_wb": 37.75,  # °C (CTI)
        "fan_power": 13.41,  # kW
        "Ka_over_L": 0.982,  # Total Ka/L
        "static_pressure_mmWG": 20.225,  # mm WG
        "static_pressure_Pa": 20.225 * 9.81,  # Pa
        "water_loading": 8.95  # l/s·m² = 32.22 m³/h·m²
    }
    
    # Compare
    comparison = {
        "your_calculation": {
            "T_cold": results["T_cold_achieved"],
            "approach": results["approach"],
            "fan_power": results["fan_power"],
            "Ka_over_L": results["Ka_over_L"],
            "static_pressure": results["total_static_pressure"],
            "water_loading": results["water_loading"],
            "K_value": results["K_value"]  # Added K value
        },
        "supplier_claimed": supplier_claimed,
        "differences": {
            "T_cold_diff": results["T_cold_achieved"] - supplier_claimed["T_cold_achieved"],
            "fan_power_diff": results["fan_power"] - supplier_claimed["fan_power"],
            "Ka_over_L_diff": results["Ka_over_L"] - supplier_claimed["Ka_over_L"]
        }
    }
    
    return results, comparison

# ============================================================================
# REPORT GENERATION
# ============================================================================

def generate_txt_report(design_results):
    """Generate a detailed TXT report"""
    report = []
    report.append("=" * 70)
    report.append("COOLING TOWER DESIGN REPORT")
    report.append("=" * 70)
    report.append(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append(f"Fill: {design_results['fill_name']}")
    report.append(f"Tower Type: {design_results['tower_name']}")
    if design_results.get('use_pdf_data', False):
        report.append(f"Data Source: PDF Performance Graphs")
    report.append("=" * 70)
    
    # Design Inputs
    report.append("\nDESIGN INPUTS")
    report.append("-" * 40)
    report.append(f"Water Flow Rate: {design_results['L']:.2f} kg/s")
    report.append(f"Air Flow Rate: {design_results['G']:.2f} kg/s")
    report.append(f"L/G Ratio: {design_results['L_over_G']:.3f}")
    report.append(f"Hot Water In: {design_results['T_hot']:.1f} °C")
    report.append(f"Target Cold Water Out: {design_results['T_cold_target']:.1f} °C")
    report.append(f"Ambient Wet Bulb: {design_results['Twb']:.1f} °C")
    report.append(f"Ambient Dry Bulb: {design_results['Tdb']:.1f} °C")
    report.append(f"Relative Humidity: {design_results['RH']:.1f} %")
    report.append(f"Site Altitude: {design_results['altitude']} m ASL")
    report.append(f"Tower Type: {design_results['tower_name']}")

    geom = design_results.get('geometry_input', {})
    if geom:
        report.append(f"Tower Shape: {geom.get('tower_shape', 'Not specified')}")
        if geom.get('tower_shape') == 'Rectangle':
            report.append(f"Input Fill Face Length: {geom.get('fill_length_m', 0):.3f} m")
            report.append(f"Input Fill Face Width/Breadth: {geom.get('fill_width_m', 0):.3f} m")
        elif geom.get('tower_shape') == 'Round':
            report.append(f"Input Tower Diameter: {geom.get('diameter_m', 0):.3f} m")
        elif 'Crossflow' in str(geom.get('tower_shape', '')):
            report.append(f"Crossflow Geometry Basis: {geom.get('tower_shape')}")

    if design_results.get('tower_type') == 'crossflow':
        cfg = design_results.get('crossflow_geometry', {})
        if cfg:
            report.append(f"Crossflow Fill Height H: {cfg.get('crossflow_fill_height_m', design_results['fill_depth']):.3f} m")
            report.append(f"Crossflow Air Travel Depth AT: {cfg.get('crossflow_air_travel_depth_m', 0):.3f} m")
            report.append(f"Crossflow Stack Width per Bank: {cfg.get('crossflow_stack_width_m', 0):.3f} m")
            bank_count_report = int(cfg.get('crossflow_bank_count', 1))
            report.append(f"Crossflow Fill Banks: {bank_count_report} ({'dual-bank typical induced-draft' if bank_count_report == 2 else 'single-bank'})")
        else:
            report.append(f"Crossflow Fill Height H: {design_results['fill_depth']:.3f} m")
        report.append(f"Total Air Face Area (banks × H × Width): {design_results['face_area']:.2f} m²")
        report.append(f"Total Water Plan Area (banks × AT × Width): {design_results.get('water_plan_area', 0):.2f} m²")
        if str(design_results.get('fill_name', '')).startswith('Brentwood XF75'):
            report.append("XF75 catalogue: D = air travel depth, W = module/stacking width, L = vertical fill height")
            report.append(f"XF75 product surface area: {XF75_CATALOG_SURFACE_AREA_M2_M3:.1f} m²/m³")
            report.append(f"XF75 sheet spacing: {XF75_CATALOG_SHEET_SPACING_MM:.1f} mm; sheets per foot: {XF75_CATALOG_SHEETS_PER_FOOT}")
    else:
        report.append(f"Fill Depth: {design_results['fill_depth']:.3f} m")
        report.append(f"Air Face Area: {design_results['face_area']:.2f} m²")
    
    # Design Results
    report.append("\nDESIGN RESULTS")
    report.append("-" * 40)
    report.append(f"Achieved Cold Water: {design_results['T_cold_achieved']:.2f} °C")
    report.append(f"Required Heat Rejection: {design_results.get('Q_target', 0):.0f} kW")
    report.append(f"Achievable Heat Rejection: {design_results['Q_achieved']:.0f} kW")
    report.append(f"Heat Rejection Margin: {design_results['Q_achieved'] - design_results.get('Q_target', 0):.0f} kW")
    report.append(f"Cooling Range: {design_results['cooling_range']:.2f} °C")
    report.append(f"Approach: {design_results['approach']:.2f} °C")
    report.append(f"NTU: {design_results['NTU']:.3f}")
    report.append(f"Ka/L: {design_results['Ka_over_L']:.3f}")
    report.append(f"KaH/L or KaV/L: {design_results.get('KaH_over_L', design_results['Ka_over_L'] * design_results['fill_depth']):.3f}")
    report.append(f"K Value (Ka): {design_results['K_value']:.3f} kW/°C")  # Added K value
    
    # Geometry and Hydraulics
    report.append("\nGEOMETRY & HYDRAULICS")
    report.append("-" * 40)
    report.append(f"Fill Volume: {design_results['fill_volume']:.2f} m³")
    report.append(f"Total Surface Area: {design_results['total_surface_area']:.0f} m²")
    report.append(f"Water Loading: {design_results['water_loading']:.1f} m³/h·m²")
    report.append(f"Water Velocity in Channels: {design_results['water_velocity']:.3f} m/s")
    report.append(f"Water Film Thickness: {design_results['water_film_thickness']} mm")
    report.append(f"Air Face Velocity: {design_results['air_face_velocity']:.2f} m/s")
    report.append(f"Net Open Air Area: {design_results.get('open_air_area', 0):.2f} m²")
    report.append(f"Air Velocity Through Fill Open Area: {design_results.get('air_velocity_through_fill', design_results['air_face_velocity']):.2f} m/s")
    report.append(f"Air Density: {design_results['air_density']:.3f} kg/m³")
    report.append(f"Fan Airflow: {design_results['air_flow_volumetric']:.2f} m³/s")
    report.append(f"Pressure Drop Method: {design_results.get('pressure_drop_method', 'Not specified')}")
    if design_results.get('uses_published_dp_equation', False):
        report.append(f"Published Fill Pressure Drop Equation Result: {design_results.get('delta_P_base', 0):.1f} Pa")
    else:
        report.append(f"Base Pressure Drop: {design_results.get('delta_P_base', 0):.1f} Pa/m (legacy curve basis)")
    report.append(f"Fill Pressure Drop: {design_results['fill_pressure_drop']:.1f} Pa")
    report.append(f"Other/Non-Fill Losses: {design_results.get('other_losses_total', 0):.1f} Pa")
    for loss_name, loss_value in design_results.get('additional_losses', {}).items():
        report.append(f"  - {loss_name}: {loss_value:.1f} Pa")
    report.append(f"Total Static Pressure: {design_results['total_static_pressure']:.1f} Pa")
    report.append(f"Estimated Fan Power: {design_results['fan_power']:.2f} kW")
    
    # Tower Characteristics
    report.append("\nTOWER CHARACTERISTICS")
    report.append("-" * 40)
    report.append(f"Type: {design_results['tower_name']}")
    report.append(f"Fill Utilization Factor: {design_results['tower_efficiency_factor']:.2f}")
    report.append(f"Pressure Drop Factor Applied to Fill ΔP: {design_results.get('tower_pressure_factor_applied', 1.0):.1f}")
    report.append(f"Legacy/Typical Tower Pressure Factor Reference: {design_results.get('legacy_tower_pressure_factor', TOWER_TYPES[design_results['tower_type']]['typical_pressure_drop_factor']):.1f}")
    
    # Fill Characteristics
    report.append("\nFILL CHARACTERISTICS")
    report.append("-" * 40)
    report.append(f"Surface Area Density: {design_results['surface_area_density']} m²/m³")
    report.append(f"Hydraulic Diameter: {design_results['hydraulic_diameter']:.1f} mm")
    report.append(f"Flute Angle: {design_results['flute_angle']}°")
    report.append(f"Free Area Fraction: {design_results['free_area_fraction']:.2f}")
    report.append(f"Free Area Source: {design_results.get('free_area_source', 'Engineering estimate - not directly published in available brochure')}")
    report.append("Free Area Note: Brentwood fill brochures generally publish surface area/sheet spacing/flute details; net open-area fraction is not directly published in the available brochure, so this value is an assumption unless manufacturer data is provided.")
    report.append(f"Fouling Risk: {design_results['fouling_risk']['risk_level']}")

    # Brentwood XF75 SI graph validity checks
    if design_results.get('xf75_graph_checks'):
        report.append("\nBRENTWOOD XF75 SI GRAPH VALIDITY CHECKS")
        report.append("-" * 40)
        if design_results['xf75_graph_checks'].get('all_ok'):
            report.append("✅ All checked operating points are within the Brentwood XF75 SI graph envelopes.")
        else:
            report.append("⚠️ One or more operating points are outside the Brentwood XF75 SI graph envelopes.")
        for label, ok, message in design_results['xf75_graph_checks'].get('checks', []):
            report.append(("✅ " if ok else "⚠️ ") + f"{label}: {message}")
    
    # Status
    report.append("\nPERFORMANCE STATUS")
    report.append("-" * 40)
    if design_results['T_cold_achieved'] <= design_results['T_cold_target']:
        report.append("✅ DESIGN MEETS REQUIREMENTS")
    else:
        report.append("⚠️ DESIGN DOES NOT MEET REQUIREMENTS")
        report.append(f"   Required improvement: {design_results['T_cold_achieved'] - design_results['T_cold_target']:.2f} °C")
    
    # Warnings
    if design_results['operating_warnings']:
        report.append("\nOPERATING WARNINGS")
        report.append("-" * 40)
        for warning in design_results['operating_warnings']:
            report.append(f"⚠️ {warning}")
    
    report.append("\n" + "=" * 70)
    report.append("END OF REPORT")
    report.append("=" * 70)
    
    return "\n".join(report)

def generate_pdf_report_from_text(report_text: str, title: str = "Cooling Tower Report") -> bytes:
    """Create a simple A4 PDF from a plain-text report for easy printing."""
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # Margins
    left = 18 * mm
    right = 18 * mm
    top = 18 * mm
    bottom = 18 * mm

    # Typography
    font = "Helvetica"
    font_bold = "Helvetica-Bold"
    font_size = 10
    line_h = 12  # points

    # Title
    y = height - top
    c.setFont(font_bold, 14)
    c.drawString(left, y, title)
    y -= 18

    c.setFont(font, font_size)

    # Wrap long lines so they stay within margins
    max_width = width - left - right

    def wrap_line(s: str):
        # crude but reliable wrap using stringWidth
        words = s.split(" ")
        out, cur = [], ""
        for w in words:
            trial = (cur + " " + w).strip()
            if c.stringWidth(trial, font, font_size) <= max_width:
                cur = trial
            else:
                if cur:
                    out.append(cur)
                cur = w
        if cur:
            out.append(cur)
        return out if out else [""]

    for raw_line in report_text.splitlines():
        # page break if needed
        if y <= bottom:
            c.showPage()
            y = height - top
            c.setFont(font, font_size)

        # blank line handling
        if raw_line.strip() == "":
            y -= line_h
            continue

        for wrapped in wrap_line(raw_line):
            if y <= bottom:
                c.showPage()
                y = height - top
                c.setFont(font, font_size)
            c.drawString(left, y, wrapped)
            y -= line_h

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer.getvalue()


# ============================================================================
# VENDOR QUOTE ANALYSIS + SPEC LIBRARY HELPERS
# ============================================================================

LIBRARY_PATH = Path(os.environ.get("COOLING_TOWER_LIBRARY_PATH", "vendor_specs_library.csv"))

def generate_docx_report(title: str, sections: list) -> bytes:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    doc.add_heading(title, 0)
    for heading, content in sections:
        doc.add_heading(str(heading), level=1)
        if isinstance(content, pd.DataFrame):
            if content.empty:
                doc.add_paragraph('No data available.')
            else:
                table = doc.add_table(rows=1, cols=len(content.columns))
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                for i, col in enumerate(content.columns):
                    hdr[i].text = str(col)
                for _, row in content.iterrows():
                    cells = table.add_row().cells
                    for i, col in enumerate(content.columns):
                        cells[i].text = str(row[col])
        elif isinstance(content, dict):
            for k, v in content.items():
                doc.add_paragraph(f"{k}: {v}")
        elif isinstance(content, (list, tuple)):
            for item in content:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            for line in str(content).splitlines():
                doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def extract_text_from_uploaded_file(uploaded_file):
    try:
        name = (uploaded_file.name or '').lower()
        if name.endswith('.pdf'):
            text_parts = []
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ''
                    if t:
                        text_parts.append(t)
            uploaded_file.seek(0)
            return '\n'.join(text_parts)
        elif name.endswith('.docx'):
            document = docx.Document(uploaded_file)
            text = '\n'.join([p.text for p in document.paragraphs if p.text])
            uploaded_file.seek(0)
            return text
        else:
            raw = uploaded_file.read()
            uploaded_file.seek(0)
            if isinstance(raw, bytes):
                for enc in ('utf-8', 'utf-16', 'latin-1'):
                    try:
                        return raw.decode(enc)
                    except Exception:
                        pass
                return raw.decode('utf-8', errors='ignore')
            return str(raw)
    except Exception:
        try:
            uploaded_file.seek(0)
        except Exception:
            pass
        return ''

def extract_text_from_multiple_files(uploaded_files):
    texts = []
    for uf in uploaded_files or []:
        tx = extract_text_from_uploaded_file(uf)
        if tx:
            texts.append(f"\n\n===== FILE: {uf.name} =====\n" + tx)
    return '\n'.join(texts)

def _to_float(num_str):
    if num_str is None:
        return None
    s = str(num_str).strip().replace(',', '')
    try:
        return float(s)
    except Exception:
        return None

def convert_temperature_to_c(value, unit):
    if value is None:
        return None
    unit = (unit or 'C').replace('°', '').upper()
    return (value - 32.0) * 5.0 / 9.0 if unit == 'F' else value

def convert_water_flow_to_m3h(value, unit):
    if value is None:
        return None
    u = (unit or 'm3/h').lower().replace('³', '3')
    if u == 'm3/h': return value
    if u in ('l/s', 'lps'): return value * 3.6
    if u in ('lpm', 'l/min'): return value * 0.06
    if u in ('usgpm', 'gpm'): return value * 0.2271247
    return value

def convert_airflow_to_m3s(value, unit):
    if value is None:
        return None
    u = (unit or 'm3/s').lower().replace('³', '3')
    if u == 'm3/s': return value
    if u == 'm3/h': return value / 3600.0
    if u == 'cfm': return value * 0.00047194745
    return value

def convert_pressure_to_pa(value, unit):
    if value is None:
        return None
    u = (unit or 'pa').lower()
    if u == 'pa': return value
    if u in ('mmaq', 'mmwg', 'mmh2o'): return value * 9.80665
    if u in ('in.wg', 'inwg', 'inh2o'): return value * 249.0889
    if u == 'kpa': return value * 1000.0
    return value

def convert_heat_to_kw(value, unit):
    if value is None:
        return None
    u = (unit or 'kw').lower()
    if u == 'kw': return value
    if u in ('kcal/h', 'kcalhr'): return value * 0.001163
    if u == 'tr': return value * 3.517
    if u in ('btu/h', 'btuhr'): return value * 0.00029307107
    return value

def convert_length_to_m(value, unit):
    if value is None:
        return None
    u = (unit or 'm').lower()
    if u == 'm': return value
    if u == 'mm': return value / 1000.0
    if u == 'cm': return value / 100.0
    if u in ('ft', 'feet'): return value * 0.3048
    if u in ('in', 'inch', 'inches'): return value * 0.0254
    return value

def parse_vendor_quote_text(text):
    cleaned = (text or '').replace('×', 'x')
    full = '\n'.join([ln.strip() for ln in cleaned.splitlines() if ln.strip()])
    low = full.lower()
    out = {
        'tower_type': None, 'water_flow_m3h': None, 'hot_water_c': None, 'cold_water_c': None,
        'wet_bulb_c': None, 'dry_bulb_c': None, 'air_flow_m3s': None, 'fan_power_kw': None,
        'static_pressure_pa': None, 'heat_load_kw': None, 'g_over_l': None, 'l_over_g': None,
        'evaporation_loss_percent': None, 'drift_loss_percent': None, 'blowdown_loss_percent': None,
        'makeup_water_percent': None, 'length_m': None, 'width_m': None, 'height_m': None, 'source_summary': []
    }
    if 'counter flow' in low or 'counterflow' in low:
        out['tower_type'] = 'counterflow_induced'
    elif 'cross flow' in low or 'crossflow' in low:
        out['tower_type'] = 'crossflow'

    def s(patterns, flags=re.I):
        for pat in patterns:
            m = re.search(pat, full, flags)
            if m:
                return m
        return None

    # water flow
    m = s([
        r'flow\s*[:\-]?\s*([\d,.]+)\s*(m3/h|m³/h|LPS|L/s|LPM|USGPM|GPM)',
        r'cooling\s*water\s*flow\s*rate\s*[:\-]?\s*([\d,.]+)\s*(m3/h|m³/h|LPS|L/s|LPM|USGPM|GPM)',
        r'design\s*flow\s*rate\s*[:\-]?\s*([\d,.]+)\s*(LPS|L/s|m3/h|m³/h|USGPM|GPM)',
        r'water\s*flow\s*rate\s*[:\-]?\s*([\d,.]+)\s*(m3/h|m³/h|LPS|L/s|LPM|USGPM|GPM)',
        r'performance\s*data\s*water\s*flow\s*rate\s*([\d,.]+)\s*(USGPM|GPM|m3/h|m³/h|LPS|L/s)'
    ])
    if m:
        out['water_flow_m3h'] = convert_water_flow_to_m3h(_to_float(m.group(1)), m.group(2))
        out['source_summary'].append(f"Water flow: {m.group(0)}")

    # hot/cold as pair or individual
    pair_patterns = [
        r'entering\s*fluid\s*temp\s*[:\-]?\s*([\d.]+)\s*[°]?(C|F).*?leaving\s*fluid\s*temp\s*[:\-]?\s*([\d.]+)\s*[°]?(C|F)',
        r'inlet\s*water\s*temperature\s*/\s*outlet\s*water\s*temperature\s*[:\-]?\s*([\d\.]+)\s*[°]?(C|F)\s*/\s*([\d\.]+)\s*[°]?(C|F)',
        r'inlet\s*temperature\s*[:\-]?\s*([\d\.]+)\s*[°]?(C|F).*?outlet\s*temperature\s*[:\-]?\s*([\d\.]+)\s*[°]?(C|F)',
        r'inlet\s*temperature\s*\(°C\s*/\s*°F\)\s*([\d\.]+)\s*(C|F)',
    ]
    m = s(pair_patterns, flags=re.I|re.S)
    if m and len(m.groups()) >= 4:
        out['hot_water_c'] = convert_temperature_to_c(_to_float(m.group(1)), m.group(2))
        out['cold_water_c'] = convert_temperature_to_c(_to_float(m.group(3)), m.group(4))
        out['source_summary'].append(f"Inlet/outlet pair: {m.group(0)}")
    # Genius style separate lines
    mh = s([r'inlet\s*temperature.*?([\d\.]+)\s*[°]?(C|F)', r'hot\s*water.*?([\d\.]+)\s*[°]?(C|F)', r'entering\s*fluid\s*temp.*?([\d\.]+)\s*[°]?(C|F)'])
    mc = s([r'outlet\s*temperature.*?([\d\.]+)\s*[°]?(C|F)', r'cold\s*water.*?([\d\.]+)\s*[°]?(C|F)', r'leaving\s*fluid\s*temp.*?([\d\.]+)\s*[°]?(C|F)'])
    if out['hot_water_c'] is None and mh:
        out['hot_water_c'] = convert_temperature_to_c(_to_float(mh.group(1)), mh.group(2))
        out['source_summary'].append(f"Hot/inlet temp: {mh.group(0)}")
    if out['cold_water_c'] is None and mc:
        out['cold_water_c'] = convert_temperature_to_c(_to_float(mc.group(1)), mc.group(2))
        out['source_summary'].append(f"Cold/outlet temp: {mc.group(0)}")

    # WB/DB
    m = s([
        r'ambient\s*dry\s*bulb\s*temperature\s*/\s*wet\s*bulb\s*temperature\s*[:\-]?\s*([\d\.]+)\s*[°]?(C|F)\s*/\s*([\d\.]+)\s*[°]?(C|F)',
        r'dry\s*bulb\s*temperature\s*/\s*wet\s*bulb\s*temperature\s*[:\-]?\s*([\d\.]+)\s*[°]?(C|F)\s*/\s*([\d\.]+)\s*[°]?(C|F)'
    ], flags=re.I|re.S)
    if m:
        out['dry_bulb_c'] = convert_temperature_to_c(_to_float(m.group(1)), m.group(2))
        out['wet_bulb_c'] = convert_temperature_to_c(_to_float(m.group(3)), m.group(4))
        out['source_summary'].append(f"DB/WB pair: {m.group(0)}")
    mw = s([r'wet\s*bulb\s*temperature.*?([\d\.]+)\s*[°]?(C|F)', r'ambient\s*wet\s*bulb.*?([\d\.]+)\s*[°]?(C|F)', r'wet\s*bulb.*?([\d\.]+)\s*[°]?(C|F)'])
    md = s([r'dry\s*bulb\s*temperature.*?([\d\.]+)\s*[°]?(C|F)', r'ambient\s*dry\s*bulb.*?([\d\.]+)\s*[°]?(C|F)', r'dry\s*bulb.*?([\d\.]+)\s*[°]?(C|F)'])
    if out['wet_bulb_c'] is None and mw:
        out['wet_bulb_c'] = convert_temperature_to_c(_to_float(mw.group(1)), mw.group(2))
        out['source_summary'].append(f"Wet bulb: {mw.group(0)}")
    if out['dry_bulb_c'] is None and md:
        out['dry_bulb_c'] = convert_temperature_to_c(_to_float(md.group(1)), md.group(2))
        out['source_summary'].append(f"Dry bulb: {md.group(0)}")

    # heat
    m = s([r'cooling\s*capacity\s*[:\-]?\s*([\d,\.]+)\s*(kW|kcal/hr|kcal/h|TR|BTU/h)', r'required\s*capacity\s*([\d,\.]+)\s*(kW|kcal/hr|kcal/h|TR|BTU/h)', r'heat\s*rejection\s*capacity\s*\(?kw\)?\s*([\d,\.]+)'])
    if m:
        unit = m.group(2) if len(m.groups()) > 1 else 'kW'
        out['heat_load_kw'] = convert_heat_to_kw(_to_float(m.group(1)), unit)
        out['source_summary'].append(f"Heat load: {m.group(0)}")

    # airflow
    m = s([r'air\s*volume\s*[:\-]?\s*([\d,\.]+)\s*(m3/s|m³/s|m3/h|m³/h|CFM)\s*[x×]\s*([\d\.]+)', r'air\s*flow\s*[:\-]?\s*([\d,\.]+)\s*(m3/s|m³/s|m3/h|m³/h|CFM)\s*[x×]\s*([\d\.]+)'])
    if m:
        base = convert_airflow_to_m3s(_to_float(m.group(1)), m.group(2)); mult = _to_float(m.group(3)) or 1
        out['air_flow_m3s'] = base * mult
        out['source_summary'].append(f"Airflow totalized: {m.group(0)}")
    else:
        m = s([r'air\s*flow\s*[:\-]?\s*([\d,\.]+)\s*(m3/s|m³/s|m3/h|m³/h|CFM)', r'air\s*volume\s*[:\-]?\s*([\d,\.]+)\s*(m3/s|m³/s|m3/h|m³/h|CFM)', r'additional\s*details.*?air\s*flow\s*[:\-]?\s*([\d,\.]+)\s*(m3/s|m³/s|m3/h|m³/h|CFM)'], flags=re.I|re.S)
        if m:
            out['air_flow_m3s'] = convert_airflow_to_m3s(_to_float(m.group(1)), m.group(2))
            out['source_summary'].append(f"Airflow: {m.group(0)}")

    # fan power
    m = s([r'required\s*power\*?\s*[:\-]?\s*([\d,\.]+)\s*(kW|HP)', r'nameplate\s*power.*?([\d,\.]+)\s*(kW|HP)', r'rated\s*kw\s*x\s*qty\s*([\d,\.]+)\s*(kW|HP)?', r'motor\s*power\s*[:\-]?\s*([\d,\.]+)\s*(kW|HP)\s*(?:total\s*([\d,\.]+)\s*sets?)?'])
    if m:
        base = _to_float(m.group(1)); unit = (m.group(2) if len(m.groups()) >= 2 and m.group(2) else 'kW')
        if unit and unit.upper() == 'HP':
            base = base * 0.7457
        mult = _to_float(m.group(3)) if len(m.groups()) >= 3 and m.group(3) else 1.0
        out['fan_power_kw'] = base * mult
        out['source_summary'].append(f"Fan/motor power: {m.group(0)}")

    # static / inlet pressure drop
    m = s([r'total\s*static\s*pressure\s*[:\-]?\s*([\d,\.]+)\s*(Pa|kPa|mmAq|mmWG|in\.wg|inwg)', r'inlet\s*pressure\s*drop\s*[:\-]?\s*([\d,\.]+)\s*(Pa|kPa|mmAq|mmWG|in\.wg|inwg)', r'static\s*pressure\s*[:\-]?\s*([\d,\.]+)\s*(Pa|kPa|mmAq|mmWG|in\.wg|inwg)'])
    if m:
        out['static_pressure_pa'] = convert_pressure_to_pa(_to_float(m.group(1)), m.group(2))
        out['source_summary'].append(f"Pressure: {m.group(0)}")

    # ratios/losses
    for key, pats in {
        'evaporation_loss_percent':[r'evaporation\s*loss\s*[%: ]+([\d\.]+)', r'evaporated\s*water\s*rate\s*[:\-]?\s*([\d\.]+)\s*LPS'],
        'drift_loss_percent':[r'drift\s*loss\s*[%: ]+([\d\.]+)'],
        'blowdown_loss_percent':[r'blowdown\s*loss\s*[%: ]+([\d\.]+)'],
        'makeup_water_percent':[r'make\s*up\s*water\s*rate\s*[%: ]+([\d\.]+)']
    }.items():
        m = s(pats)
        if m:
            out[key] = _to_float(m.group(1))
            out['source_summary'].append(f"{key}: {m.group(0)}")

    m = s([r'air[-\s]*water\s*ratio\s*[:\-]?\s*([\d\.]+)\s*kg/kg', r'gas\s*water\s*ratio\s*[:\-]?\s*([\d\.]+)', r'G/L\s*[:=]?\s*([\d\.]+)', r'L/G\s*[:=]?\s*([\d\.]+)'])
    if m:
        ratio = _to_float(m.group(1)); label = m.group(0).upper()
        if 'L/G' in label:
            out['l_over_g'] = ratio; out['g_over_l'] = 1.0/ratio if ratio else None
        else:
            out['g_over_l'] = ratio; out['l_over_g'] = 1.0/ratio if ratio else None
        out['source_summary'].append(f"Air/water ratio: {m.group(0)}")

    # dimensions inline or separate
    m = s([r'overall\s*dimensions\s*\(wxlxh\)\s*[:\-]?\s*([\d,\.]+)\s*mm\s*x\s*([\d,\.]+)\s*mm\s*x\s*([\d,\.]+)\s*mm',
           r'length\s*[x×]\s*width\s*[x×]\s*height.*?([\d,\.]+)\s*[x×]\s*([\d,\.]+)\s*[x×]\s*([\d,\.]+)\s*(mm|m)',
           r'(?:equipment\s*size|length\s*×\s*width\s*×\s*height|l\s*×\s*w\s*×\s*h)\s*[:\-]?\s*([\d,\.]+)\s*[x×]\s*([\d,\.]+)\s*[x×]\s*([\d,\.]+)\s*(mm|m)',
           r'([\d,\.]+)\s*[x×]\s*([\d,\.]+)\s*[x×]\s*([\d,\.]+)\s*(mm|m)'])
    if m:
        if len(m.groups()) == 3:
            unit='mm'; w,l,h = m.group(1), m.group(2), m.group(3)
        else:
            w,l,h,unit = m.group(1), m.group(2), m.group(3), m.group(4)
        # if WxLxH given, map appropriately
        out['width_m'] = convert_length_to_m(_to_float(w), unit)
        out['length_m'] = convert_length_to_m(_to_float(l), unit)
        out['height_m'] = convert_length_to_m(_to_float(h), unit)
        out['source_summary'].append(f"Dimensions inline: {m.group(0)}")
    for field, pats in {
        'width_m':[r'width\s*\(w\)\s*[:\-]?\s*([\d,\.]+)\s*(mm|m|ft|in)', r'dimension\s*of\s*width\s*\(w\)\s*([\d,\.]+)\s*(mm|m|ft|in)'],
        'length_m':[r'length\s*\(l\)\s*[:\-]?\s*([\d,\.]+)\s*(mm|m|ft|in)', r'weight\s*length\s*\(l\)\s*([\d,\.]+)\s*(mm|m|ft|in)'],
        'height_m':[r'total\s*height\s*\(h\)\s*[:\-]?\s*([\d,\.]+)\s*(mm|m|ft|in)', r'height\s*\(h\)\s*[:\-]?\s*([\d,\.]+)\s*(mm|m|ft|in)']
    }.items():
        if out[field] is None:
            m = s(pats)
            if m:
                out[field] = convert_length_to_m(_to_float(m.group(1)), m.group(2))
                out['source_summary'].append(f"{field}: {m.group(0)}")

    return out

def back_calculate_effective_kavl(T_hot, T_cold, Twb, tower_type='counterflow_induced'):
    try:
        num = max(T_cold - Twb, 1e-6); den = max(T_hot - Twb, 1e-6)
        ratio = min(max(num/den, 1e-6), 0.999999)
        ntu_eff = -math.log(ratio)
        kavl = ntu_eff / 1.05 if str(tower_type).startswith('counterflow') else ntu_eff
        return max(kavl, 0.0), ntu_eff
    except Exception:
        return None, None

def estimate_evaporation_loss_m3h(water_flow_m3h, cooling_range_c):
    if water_flow_m3h is None or cooling_range_c is None:
        return None
    return 0.00085 * water_flow_m3h * cooling_range_c

def build_vendor_candidate_fill_list(tower_type):
    return [f for f in (['XF75','CF1900SBMA','XF125','XF3000'] if tower_type=='crossflow' else ['CT1200AT','CF1900SBMA','CF1200','XF3000']) if f in BRENTWOOD_FILLS]

def _read_library():
    if LIBRARY_PATH.exists():
        try:
            return pd.read_csv(LIBRARY_PATH)
        except Exception:
            return pd.DataFrame()
    return pd.DataFrame()

def _write_library(df):
    try:
        df.to_csv(LIBRARY_PATH, index=False)
        return True
    except Exception:
        return False

def save_spec_to_library(spec_dict):
    row = pd.DataFrame([spec_dict])
    existing = _read_library()
    combined = pd.concat([existing, row], ignore_index=True) if not existing.empty else row
    return _write_library(combined)

def find_nearest_library_specs(target_dict, top_n=5):
    df = _read_library()
    if df.empty:
        return df
    work = df.copy()
    for col in ['heat_load_kw','water_flow_m3h','hot_water_c','cold_water_c','wet_bulb_c','length_m','width_m','effective_kavl']:
        if col not in work.columns:
            work[col] = np.nan
    score = np.zeros(len(work))
    weights = {'heat_load_kw':3.0,'water_flow_m3h':3.0,'hot_water_c':1.5,'cold_water_c':1.5,'wet_bulb_c':1.2,'length_m':1.0,'width_m':1.0,'effective_kavl':2.0}
    for col, w in weights.items():
        tv = target_dict.get(col)
        if tv is None or pd.isna(tv):
            continue
        vals = pd.to_numeric(work[col], errors='coerce')
        denom = max(abs(tv), 1e-6)
        score += w * ((vals.fillna(tv) - tv).abs() / denom)
    work['match_score'] = score
    return work.sort_values('match_score').head(top_n)

def render_vendor_comparison_mode():
    st.header('📄 Vendor Quote Analysis & Automatic Fill Comparison')
    st.caption('Upload up to 5 vendor PDF/DOCX/TXT files. The app combines them, extracts technical data, converts to SI, compares Brentwood fills, and can save the captured spec to a reusable local library.')
    st.warning('Library storage in this code is local CSV-based. On Streamlit Cloud that may not persist across redeploys/restarts. For durable storage, Google Drive/Sheets or another database backend would be better.')

    with st.sidebar:
        st.header('📎 Vendor Comparison Inputs')
        uploaded_files = st.file_uploader('Upload vendor documents', type=['pdf','docx','txt'], accept_multiple_files=True, key='vendor_docs_uploader')
        default_tower_type = st.selectbox('Default tower type if parser is unsure', options=list(TOWER_TYPES.keys()), format_func=lambda x: TOWER_TYPES[x]['name'], index=1)
        default_fill_depth = st.number_input('Assumed fill depth for comparison (m)', min_value=0.300, max_value=2.500, value=1.000, step=0.050, format='%.3f')
        default_altitude = st.number_input('Site altitude (m)', min_value=0, max_value=5000, value=0, step=50)
        use_pdf_for_ct1200 = st.checkbox('Use Brentwood CT1200AT published equations', value=True)
        use_pdf_for_xf75 = st.checkbox('Use Brentwood XF75 published equations', value=True)
        run_compare = st.button('🚀 Run Vendor Comparison', type='primary', use_container_width=True)

    extracted = {}
    combined_text = ''
    cached_vendor_files = get_cached_vendor_files()
    if cached_vendor_files:
        st.caption(f"Files ready in session cache: **{len(cached_vendor_files)}** → " + ', '.join([x.get('name','?') for x in cached_vendor_files]))
    if uploaded_files:
        if len(uploaded_files) > 5:
            st.error('Please upload a maximum of 5 files at one time.')
            return
        cache_uploaded_vendor_files(uploaded_files)
        combined_text = extract_text_from_multiple_files(uploaded_files)
        extracted = parse_vendor_quote_text(combined_text)
        st.subheader('🔎 Extracted Technical Data (best effort)')
        preview = pd.DataFrame({'Parameter': list(extracted.keys()), 'Value': [extracted[k] for k in extracted.keys()]})
        st.dataframe(preview, use_container_width=True, height=350)
        if extracted.get('source_summary'):
            with st.expander('Show extraction notes'):
                for item in extracted['source_summary']:
                    st.write('-', item)
        cap_docx = generate_docx_report('Captured Vendor Technical Data', [('Uploaded Files', [uf.name for uf in uploaded_files]), ('Extracted Data', {k:v for k,v in extracted.items() if k!='source_summary'}), ('Extraction Notes', extracted.get('source_summary', []))])
        st.download_button('📥 Download Captured Data as Word (.docx)', data=cap_docx, file_name=f'captured_vendor_data_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', use_container_width=True)

    st.subheader('✍️ Confirm / Edit Technical Inputs')
    col1, col2, col3 = st.columns(3)
    with col1:
        tower_type = st.selectbox('Tower type', options=list(TOWER_TYPES.keys()), format_func=lambda x: TOWER_TYPES[x]['name'], index=list(TOWER_TYPES.keys()).index(extracted.get('tower_type')) if extracted.get('tower_type') in TOWER_TYPES else list(TOWER_TYPES.keys()).index(default_tower_type), key='vendor_tower_type')
        water_flow_m3h = st.number_input('Water flow (m³/h)', min_value=0.0, value=float(extracted.get('water_flow_m3h') or 0.0), step=10.0, format='%.3f')
        hot_water_c = st.number_input('Hot water in (°C)', min_value=-20.0, max_value=120.0, value=float(extracted.get('hot_water_c') or 40.0), step=0.1, format='%.2f')
        cold_water_c = st.number_input('Cold water out (°C)', min_value=-20.0, max_value=120.0, value=float(extracted.get('cold_water_c') or 35.0), step=0.1, format='%.2f')
    with col2:
        wet_bulb_c = st.number_input('Wet bulb (°C)', min_value=-20.0, max_value=80.0, value=float(extracted.get('wet_bulb_c') or 32.0), step=0.1, format='%.2f')
        dry_bulb_c = st.number_input('Dry bulb (°C)', min_value=-20.0, max_value=80.0, value=float(extracted.get('dry_bulb_c') or max((extracted.get('wet_bulb_c') or 32.0) + 3.0, 35.0)), step=0.1, format='%.2f')
        air_flow_m3s = st.number_input('Vendor air flow (m³/s)', min_value=0.0, value=float(extracted.get('air_flow_m3s') or 0.0), step=1.0, format='%.3f')
        fan_power_kw_vendor = st.number_input('Vendor fan power total (kW)', min_value=0.0, value=float(extracted.get('fan_power_kw') or 0.0), step=1.0, format='%.2f')
    with col3:
        static_pressure_pa_vendor = st.number_input('Vendor static pressure (Pa)', min_value=0.0, value=float(extracted.get('static_pressure_pa') or 0.0), step=1.0, format='%.1f')
        length_m = st.number_input('Tower length (m)', min_value=0.0, value=float(extracted.get('length_m') or 0.0), step=0.1, format='%.3f')
        width_m = st.number_input('Tower width (m)', min_value=0.0, value=float(extracted.get('width_m') or 0.0), step=0.1, format='%.3f')
        fill_depth_m = st.number_input('Comparison fill depth (m)', min_value=0.300, max_value=2.500, value=float(default_fill_depth), step=0.050, format='%.3f')
    face_area = (length_m * width_m) if (length_m and width_m) else 0.0
    if face_area > 0:
        st.info(f'Using face area from vendor dimensions: **{face_area:.2f} m²**')
    else:
        face_area = st.number_input('Fallback face area (m²)', min_value=0.1, value=36.0, step=1.0, format='%.2f')
    l_over_g_guess = st.number_input('Fallback L/G ratio (used if vendor airflow missing)', min_value=0.3, max_value=5.0, value=float(extracted.get('l_over_g') or 1.25), step=0.05, format='%.3f')

    if run_compare:
        errors = []
        if water_flow_m3h <= 0: errors.append('Water flow must be > 0')
        if hot_water_c <= cold_water_c: errors.append('Hot water temperature must be greater than cold water temperature')
        if wet_bulb_c >= hot_water_c: errors.append('Wet bulb should be less than hot water temperature for a sensible comparison')
        if face_area <= 0: errors.append('Face area must be > 0')
        if errors:
            [st.error(e) for e in errors]
            return
        water_mass_kg_s = water_flow_m3h * 1000.0 / 3600.0
        rho_air_vendor = air_density_calc(dry_bulb_c, wet_bulb_c, default_altitude)
        if air_flow_m3s > 0:
            air_mass_kg_s = air_flow_m3s * rho_air_vendor
            l_over_g_used = water_mass_kg_s / air_mass_kg_s if air_mass_kg_s > 0 else None
        else:
            l_over_g_used = l_over_g_guess
            air_mass_kg_s = water_mass_kg_s / l_over_g_used
            air_flow_m3s = air_mass_kg_s / rho_air_vendor if rho_air_vendor > 0 else 0.0
        heat_kw = water_mass_kg_s * 4.186 * (hot_water_c - cold_water_c)
        effective_kavl, effective_ntu = back_calculate_effective_kavl(hot_water_c, cold_water_c, wet_bulb_c, tower_type)
        water_loading = water_flow_m3h / face_area
        air_velocity = air_flow_m3s / face_area
        evap_loss_m3h = estimate_evaporation_loss_m3h(water_flow_m3h, hot_water_c - cold_water_c)
        summary_cols = st.columns(6)
        metrics = [(summary_cols[0],'Heat Rejection',f'{heat_kw:,.0f} kW'),(summary_cols[1],'Effective KaV/L',f'{effective_kavl:.3f}' if effective_kavl is not None else '—'),(summary_cols[2],'Water Loading',f'{water_loading:.1f} m³/h·m²'),(summary_cols[3],'Air Velocity',f'{air_velocity:.2f} m/s'),(summary_cols[4],'L/G',f'{l_over_g_used:.3f}'),(summary_cols[5],'Estimated Evaporation',f'{evap_loss_m3h:.2f} m³/h' if evap_loss_m3h is not None else '—')]
        for c,lbl,val in metrics: c.metric(lbl,val)
        rows = []
        for fill in build_vendor_candidate_fill_list(tower_type):
            use_pdf = (fill=='CT1200AT' and use_pdf_for_ct1200) or (fill=='XF75' and use_pdf_for_xf75)
            res = solve_cooling_tower_enhanced(water_mass_kg_s, air_mass_kg_s, hot_water_c, cold_water_c, wet_bulb_c, dry_bulb_c, fill, tower_type, fill_depth_m, face_area, default_altitude, use_pdf)
            predicted_kavl = res['Ka_over_L'] * fill_depth_m
            rows.append({'Fill': res['fill_name'], 'Data Source':'PDF' if res.get('use_pdf_data', False) else 'Default', 'Vendor Cold Water (°C)': round(cold_water_c,2), 'Model Cold Water (°C)': round(res['T_cold_achieved'],2), 'ΔT Model-Vendor (°C)': round(res['T_cold_achieved']-cold_water_c,2), 'Required Heat (kW)': round(res.get('Q_target', heat_kw),1), 'Achievable Heat (kW)': round(res['Q_achieved'],1), 'Vendor Eff. KaV/L': round(effective_kavl,3) if effective_kavl is not None else None, 'Predicted KaV/L': round(predicted_kavl,3), 'Utilization Ratio': round((effective_kavl/predicted_kavl),3) if effective_kavl is not None and predicted_kavl>0 else None, 'Model Fan Power (kW)': round(res['fan_power'],2), 'Vendor Fan Power (kW)': round(fan_power_kw_vendor,2) if fan_power_kw_vendor else None, 'Water Loading (m³/h·m²)': round(res['water_loading'],2), 'Air Face Velocity (m/s)': round(res['air_face_velocity'],2), 'Net Open Area (m²)': round(res.get('open_air_area',0),2), 'Air Velocity Through Fill (m/s)': round(res.get('air_velocity_through_fill', res['air_face_velocity']),2), 'Free Area Fraction': round(res.get('free_area_fraction',0),2), 'Fill ΔP (Pa)': round(res['fill_pressure_drop'],1), 'Other Losses (Pa)': round(res.get('other_losses_total',0),1), 'Static Pressure (Pa)': round(res['total_static_pressure'],1)})
        df = pd.DataFrame(rows)
        if not df.empty:
            df['ABS ΔT'] = df['ΔT Model-Vendor (°C)'].abs(); df = df.sort_values(['ABS ΔT','Model Fan Power (kW)']).reset_index(drop=True)
        st.subheader('🤖 Automatic Fill Comparison')
        st.dataframe(df, use_container_width=True)
        if not df.empty:
            best = df.iloc[0]
            st.success(f"Best thermal match: **{best['Fill']}** with model cold water **{best['Model Cold Water (°C)']} °C** and ΔT **{best['ΔT Model-Vendor (°C)']} °C**.")
        csv_bytes = df.to_csv(index=False).encode('utf-8')
        st.download_button('📥 Download Vendor Comparison CSV', data=csv_bytes, file_name=f'vendor_fill_comparison_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.csv', mime='text/csv', use_container_width=True)
        vendor_data = {'tower_type': tower_type,'water_flow_m3h': round(water_flow_m3h,3),'hot_water_c': round(hot_water_c,3),'cold_water_c': round(cold_water_c,3),'wet_bulb_c': round(wet_bulb_c,3),'dry_bulb_c': round(dry_bulb_c,3),'face_area_m2': round(face_area,3),'fill_depth_m': round(fill_depth_m,3),'air_flow_m3s': round(air_flow_m3s,3),'air_mass_flow_kg_s': round(air_mass_kg_s,3),'L_over_G': round(l_over_g_used,4),'heat_rejection_kw': round(heat_kw,2),'effective_kavl': round(effective_kavl,4) if effective_kavl is not None else None,'effective_ntu': round(effective_ntu,4) if effective_ntu is not None else None,'estimated_evaporation_m3h': round(evap_loss_m3h,3) if evap_loss_m3h is not None else None,'static_pressure_pa_vendor': round(static_pressure_pa_vendor,2) if static_pressure_pa_vendor else None,'fan_power_kw_vendor': round(fan_power_kw_vendor,2) if fan_power_kw_vendor else None,'source_summary': extracted.get('source_summary', []) if extracted else []}
        report_text = '\n'.join(['='*72,'VENDOR QUOTE COMPARISON REPORT','='*72,f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",'','EXTRACTED / CONFIRMED VENDOR DATA','-'*72]+[f"{k}: {v}" for k,v in vendor_data.items() if k!='source_summary']+['','EXTRACTION NOTES','-'*72]+[f"- {x}" for x in vendor_data.get('source_summary',[])] + ['','FILL COMPARISON','-'*72, df.to_string(index=False), '', '='*72])
        st.download_button('📥 Download Vendor Comparison TXT', data=report_text, file_name=f'vendor_fill_comparison_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.txt', mime='text/plain', use_container_width=True)
        pdf_bytes = generate_pdf_report_from_text(report_text, title='Vendor Quote Comparison Report')
        st.download_button('📥 Download Vendor Comparison PDF', data=pdf_bytes, file_name=f'vendor_fill_comparison_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf', mime='application/pdf', use_container_width=True)
        docx_bytes = generate_docx_report('Vendor Quote Comparison Report', [('Confirmed Vendor Inputs', vendor_data), ('Fill Comparison', df)])
        st.download_button('📥 Download Vendor Comparison Word (.docx)', data=docx_bytes, file_name=f'vendor_fill_comparison_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', use_container_width=True)
        save_row = {'saved_at': datetime.datetime.now().isoformat(timespec='seconds'), 'source_files': ', '.join([uf.name for uf in (uploaded_files or [])]), 'tower_type': tower_type, 'water_flow_m3h': water_flow_m3h, 'hot_water_c': hot_water_c, 'cold_water_c': cold_water_c, 'wet_bulb_c': wet_bulb_c, 'dry_bulb_c': dry_bulb_c, 'length_m': length_m, 'width_m': width_m, 'height_m': extracted.get('height_m'), 'heat_load_kw': heat_kw, 'effective_kavl': effective_kavl, 'air_flow_m3s': air_flow_m3s, 'vendor_fan_power_kw': fan_power_kw_vendor, 'vendor_static_pressure_pa': static_pressure_pa_vendor}
        if st.button('💾 Save this captured spec to local library'):
            ok = save_spec_to_library(save_row)
            st.success('Saved to local library CSV.' if ok else 'Could not save to local library.')
        with st.expander('Preview vendor comparison report'):
            st.text(report_text[:4000] + ('...' if len(report_text) > 4000 else ''))

def render_nearest_specs_panel(target):
    nearest = find_nearest_library_specs(target, top_n=5)
    st.subheader('📚 Nearby Saved Vendor Specs')
    if nearest.empty:
        st.caption('No saved library specs found yet.')
    else:
        show_cols = [c for c in ['source_files','tower_type','heat_load_kw','water_flow_m3h','hot_water_c','cold_water_c','wet_bulb_c','length_m','width_m','effective_kavl','match_score'] if c in nearest.columns]
        st.dataframe(nearest[show_cols], use_container_width=True)


# ============================================================================
# STREAMLIT APP - MAIN FUNCTION WITH ENHANCED UI
# ============================================================================

def main():
    # Check password first
    if not check_password():
        st.stop()
    
    # Set page config
    st.set_page_config(
        page_title="Professional Cooling Tower Design with Enhanced UI",
        page_icon="🌊",
        layout="wide"
    )
    
    # Main title
    st.title("🌊 Complete Cooling Tower Design Tool")
    st.markdown("**Enhanced UI | CF1200 & CT1200AT Support | Counterflow Towers | Supplier Validation | Vendor Quote Analysis | Spec Library**")

    app_mode = st.radio("Choose Workspace", ["Cooling Tower Design", "Vendor Quote Comparison"], horizontal=True)
    if app_mode == "Vendor Quote Comparison":
        render_vendor_comparison_mode()
        return
    
    # Initialize session state for geometry
    if 'tower_shape' not in st.session_state:
        st.session_state.tower_shape = "Rectangle"
    if 'face_area' not in st.session_state:
        st.session_state.face_area = 36.94
    
    # Initialize session state for PDF data usage
    if 'use_pdf_data' not in st.session_state:
        st.session_state.use_pdf_data = False

    if 'use_xf75_pdf_data' not in st.session_state:
        st.session_state.use_xf75_pdf_data = False
    if 'xf75_air_travel_depth' not in st.session_state:
        st.session_state.xf75_air_travel_depth = 0.914
    
    # ========================================================================
    # SIDEBAR - DESIGN INPUTS WITH ENHANCED CONTROLS
    # ========================================================================
    with st.sidebar:
        st.header("📥 Design Inputs")
        
        # Calculation Mode
        calc_mode = st.radio(
            "**Calculation Mode:**",
            ["Mode 1: Given Heat Load → Find Water Flow",
             "Mode 2: Given Water Flow → Find Heat Load"],
            help="Choose whether you know the heat load or water flow rate"
        )
        
        # Temperature inputs - Using number_input with step for +/- buttons
        st.subheader("🌡️ Temperature Conditions")
        
        col1, col2 = st.columns(2)
        with col1:
            T_hot = st.number_input("Hot Water In (°C)", 
                                   value=37.0, min_value=20.0, max_value=60.0, 
                                   step=0.5, format="%.1f",
                                   help="Inlet water temperature to cooling tower")
        with col2:
            T_cold_target = st.number_input("Target Cold Water Out (°C)", 
                                           value=32.0, min_value=10.0, max_value=40.0, 
                                           step=0.5, format="%.1f",
                                           help="Desired outlet water temperature")
        
        col3, col4 = st.columns(2)
        with col3:
            Twb = st.number_input("Ambient Wet Bulb (°C)", 
                                 value=28.0, min_value=10.0, max_value=40.0, 
                                 step=0.5, format="%.1f",
                                 help="Critical design parameter for cooling towers")
        with col4:
            # ADDED: Dry bulb temperature input
            Tdb = st.number_input("Ambient Dry Bulb (°C)", 
                                 value=33.0, min_value=10.0, max_value=50.0, 
                                 step=0.5, format="%.1f",
                                 help="Ambient air dry bulb temperature")
        
        # Mode-specific inputs with number_input
        st.subheader("💧 Flow Parameters")
        if calc_mode == "Mode 1: Given Heat Load → Find Water Flow":
            Q_input = st.number_input("Heat Load to Remove (kW)", 
                                     value=2090.0, min_value=100.0, max_value=10000.0,
                                     step=100.0, format="%.1f")
            Cp = 4.186
            if T_hot > T_cold_target:
                L = Q_input / (Cp * (T_hot - T_cold_target))
            else:
                L = 100.0
                st.error("Hot water temperature must be greater than cold water target")
            st.metric("Calculated Water Flow", f"{L:.2f} kg/s")
            # Display in LPM for convenience (assumes ~1000 kg/m³)
            water_lpm = (L / 1000.0) * 60.0 * 1000.0
            st.caption(f"≈ **{water_lpm:,.0f} LPM** (SI; assuming water density 1000 kg/m³)")
        else:
            L = st.number_input("Water Flow Rate (kg/s)", 
                               value=100.0, min_value=10.0, max_value=500.0,
                               step=5.0, format="%.2f")
            Cp = 4.186
            if T_hot > T_cold_target:
                Q_input = L * Cp * (T_hot - T_cold_target)
            else:
                Q_input = 2090.0
                st.error("Hot water temperature must be greater than cold water target")
            st.metric("Calculated Heat Load", f"{Q_input:.0f} kW")
        
        # Air Flow Specification with L/G as number_input
        st.subheader("🌬️ Air Flow Specification")
        
        air_input_method = st.radio(
            "Air Flow Input Method:",
            ["Method 1: Set L/G Ratio",
             "Method 2: Set Direct Air Flow Rate"],
            help="Choose to set L/G ratio or direct air flow rate"
        )
        
        if air_input_method == "Method 1: Set L/G Ratio":
            # CHANGED: Using number_input for L/G ratio
            L_over_G = st.number_input("L/G Ratio (Liquid to Gas mass ratio)", 
                                      value=1.25, min_value=0.5, max_value=3.0,
                                      step=0.05, format="%.3f",
                                      help="Typical range: 0.8-1.5. Higher = more air, lower pressure drop")
            G = L / L_over_G
            st.metric("Calculated Air Flow", f"{G:.2f} kg/s")
        else:
            G = st.number_input("Air Mass Flow Rate (kg/s)", 
                               value=80.0, min_value=10.0, max_value=300.0,
                               step=5.0, format="%.2f")
            L_over_G = L / G
            st.metric("Calculated L/G Ratio", f"{L_over_G:.3f}")
        
        
        # --------------------------------------------------------------------
        # Derived Flows (SI) - convenience display (does not change calculations)
        # --------------------------------------------------------------------
        try:
            rho_air = air_density_calc(Tdb, Twb, altitude)
            air_vol_flow_m3s = G / rho_air if rho_air > 0 else 0.0
        except Exception:
            rho_air = 0.0
            air_vol_flow_m3s = 0.0

        # Water volumetric flow for display
        water_flow_m3s = L / 1000.0  # m³/s (rho≈1000 kg/m³)
        water_flow_m3h = water_flow_m3s * 3600.0
        water_flow_lpm = water_flow_m3s * 60.0 * 1000.0

        st.subheader("📌 Derived Flows (SI)")
        cdf1, cdf2 = st.columns(2)
        with cdf1:
            st.metric("Water Flow", f"{water_flow_lpm:,.0f} LPM")
            st.caption(f"({water_flow_m3h:,.1f} m³/h)")
        with cdf2:
            st.metric("Air Flow (Vol.)", f"{air_vol_flow_m3s:,.2f} m³/s")
            if rho_air > 0:
                st.caption(f"(ρ_air ≈ {rho_air:.3f} kg/m³)")

# Tower Type Selection
        st.subheader("🏗️ Tower Configuration")
        tower_type = st.selectbox(
            "Tower Type:",
            options=list(TOWER_TYPES.keys()),
            format_func=lambda x: TOWER_TYPES[x]["name"],
            help="Select tower flow arrangement and draft type"
        )
        
        # Show tower description
        tower_desc = TOWER_TYPES[tower_type]["description"]
        st.caption(f"*{tower_desc}*")
        
        # Geometry Parameters
        st.subheader("📐 Geometry Parameters")

        if tower_type == "crossflow":
            st.markdown("**Crossflow fill geometry**")
            xf75_published = st.radio(
                "Crossflow geometry input mode:",
                ["Use Brentwood XF75 published H/AT selections", "Custom crossflow fill geometry"],
                horizontal=False,
                help="For XF75, use the SI published height and air-travel-depth families. For other crossflow fills, choose custom geometry."
            )

            if xf75_published == "Use Brentwood XF75 published H/AT selections":
                fill_depth = st.selectbox(
                    "XF75 Fill Height H (m)",
                    options=XF75_SI_FILL_HEIGHT_OPTIONS_M,
                    index=0,
                    format_func=lambda x: f"{x:.3f} m ({x/0.3048:.0f} ft)",
                    help="Brentwood XF75 SI thermal charts are published at 1.829, 2.438, 3.657 and 4.877 m fill heights."
                )
                xf75_air_travel_depth = st.selectbox(
                    "XF75 Air Travel Depth AT (m)",
                    options=XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M,
                    index=1,
                    format_func=lambda x: f"{x:.3f} m ({x/0.3048:.0f} ft)",
                    help="Horizontal air path through the fill. Brentwood SI pressure-drop charts are published at 0.610 to 1.829 m AT."
                )
            else:
                fill_depth = st.number_input(
                    "Custom Crossflow Fill Height H (m)",
                    value=1.829, min_value=0.300, max_value=6.000,
                    step=0.050, format="%.3f",
                    help="Vertical fill height/water-fall height. If outside XF75 chart heights, the app will warn when XF75 equations are used."
                )
                xf75_air_travel_depth = st.number_input(
                    "Custom Air Travel Depth AT (m)",
                    value=0.914, min_value=0.300, max_value=3.000,
                    step=0.050, format="%.3f",
                    help="Horizontal air path through fill. If outside XF75 chart AT values, the app will warn when XF75 equations are used."
                )

            crossflow_stack_width = st.number_input(
                "Fill Stack Width per Bank / Tower Length (m)",
                value=6.000, min_value=0.500, max_value=50.000,
                step=0.100, format="%.3f",
                help="Width/length of each fill bank along the tower. For dual-bank crossflow, total air face area = 2 × H × width."
            )

            crossflow_bank_count_label = st.radio(
                "Crossflow Fill Bank Arrangement",
                ["Dual fill banks (typical induced-draft crossflow)", "Single fill bank"],
                index=0,
                help="Most induced-draft crossflow towers have two fill banks, one on each side of a central plenum/fan. Select single-bank only for special one-sided arrangements."
            )
            crossflow_bank_count = 2 if crossflow_bank_count_label.startswith("Dual") else 1

            air_face_area_single_bank = fill_depth * crossflow_stack_width
            water_plan_area_single_bank = xf75_air_travel_depth * crossflow_stack_width
            fill_volume_single_bank = fill_depth * xf75_air_travel_depth * crossflow_stack_width
            air_face_area = crossflow_bank_count * air_face_area_single_bank
            water_plan_area = crossflow_bank_count * water_plan_area_single_bank
            fill_volume_preview = crossflow_bank_count * fill_volume_single_bank
            xf75_surface_area_preview = fill_volume_preview * XF75_CATALOG_SURFACE_AREA_M2_M3
            xf75_module_est = xf75_catalog_module_estimate(fill_depth, xf75_air_travel_depth, crossflow_stack_width)
            face_area = air_face_area
            tower_shape = "Crossflow rectangular fill bank"
            st.session_state.geometry_input = {
                "tower_shape": tower_shape,
                "crossflow_fill_height_m": float(fill_depth),
                "crossflow_air_travel_depth_m": float(xf75_air_travel_depth),
                "crossflow_stack_width_m": float(crossflow_stack_width),
                "crossflow_bank_count": int(crossflow_bank_count),
            }

            st.session_state.crossflow_fill_height_m = float(fill_depth)
            st.session_state.crossflow_air_travel_depth_m = float(xf75_air_travel_depth)
            st.session_state.crossflow_stack_width_m = float(crossflow_stack_width)
            st.session_state.crossflow_bank_count = int(crossflow_bank_count)
            st.session_state.xf75_air_travel_depth = float(xf75_air_travel_depth)  # legacy key used elsewhere
            st.session_state.face_area = float(face_area)
            st.session_state.crossflow_water_plan_area_m2 = float(water_plan_area)
            st.session_state.crossflow_fill_volume_m3 = float(fill_volume_preview)

            st.success(
                f"Banks = {crossflow_bank_count} | Total air face area = {air_face_area:.2f} m² | "
                f"Total water plan area = {water_plan_area:.2f} m² | Total fill volume = {fill_volume_preview:.2f} m³"
            )
            if crossflow_bank_count == 2:
                st.caption(
                    f"Per bank: air face area = {air_face_area_single_bank:.2f} m², "
                    f"water plan area = {water_plan_area_single_bank:.2f} m², "
                    f"fill volume = {fill_volume_single_bank:.2f} m³."
                )
            st.caption("For crossflow, the app uses separate air-face and water-plan areas. Counterflow geometry is unchanged.")

            with st.expander("📦 XF75 catalogue geometry interpretation", expanded=False):
                st.markdown(
                    """
                    **XF75 catalogue notation:** **D = air travel depth**, **W = media/module stacking width**, and **L = vertical fill height**.
                    Catalogue surface area is **51 ft²/ft³ = 167.4 m²/m³**, sheet spacing is **0.75 in = 19 mm**, and there are **16 sheets/ft**.
                    """
                )
                cat_cols = st.columns(4)
                with cat_cols[0]:
                    st.metric("Catalogue D", "0.610 m")
                with cat_cols[1]:
                    st.metric("Catalogue W", "0.305 m")
                with cat_cols[2]:
                    st.metric("Catalogue L options", "1.829–3.658 m")
                with cat_cols[3]:
                    st.metric("Surface area", f"{xf75_surface_area_preview:,.0f} m²")
                st.caption(
                    f"Estimated packing: {xf75_module_est['depth_modules_est']} module(s) in air-depth direction × "
                    f"{xf75_module_est['width_modules_est']} module(s) across width; "
                    f"vertical layer estimate = {xf75_module_est['vertical_layers_est']}. "
                    "The 4.877 m / 16 ft thermal chart height is above the standard single catalogue pack length and normally implies vertical stacking/custom arrangement."
                )

        else:
            # Counterflow geometry remains unchanged.
            fill_depth = st.number_input("Fill Depth (m)", 
                                        value=0.600, min_value=0.300, max_value=2.000,
                                        step=0.050, format="%.3f",
                                        help="Depth of fill media in air flow direction")
            
            st.markdown("**Tower Shape Selection**")
            tower_shape = st.radio(
                "Select tower shape:",
                ["Rectangle", "Round"],
                horizontal=True,
                key="tower_shape_selector"
            )
            
            if tower_shape == "Rectangle":
                col1, col2 = st.columns(2)
                with col1:
                    fill_length = st.number_input("Fill Face Length (m)", 
                                                 value=6.08, min_value=1.0, max_value=20.0,
                                                 step=0.1, format="%.2f")
                with col2:
                    fill_width = st.number_input("Fill Face Width/Breadth (m)", 
                                                value=6.08, min_value=1.0, max_value=20.0,
                                                step=0.1, format="%.2f")
                face_area = fill_length * fill_width
                st.success(f"**Calculated Face Area:** {face_area:.2f} m²")
                
            else:  # Round tower
                diameter = st.number_input("Tower Diameter (m)", 
                                          value=6.85, min_value=1.0, max_value=20.0,
                                          step=0.1, format="%.2f")
                face_area = math.pi * (diameter / 2) ** 2
                st.success(f"**Calculated Face Area:** {face_area:.2f} m²")
            
            st.session_state.face_area = face_area

        # CHANGED: Altitude as number_input with clear label
        altitude = st.number_input("Site Altitude from Sea Level (m)", 
                                  value=0, min_value=0, max_value=3000,
                                  step=100, format="%d",
                                  help="Altitude above sea level for air density correction")
        
        # Fill Selection - filtered by tower type to prevent misuse of correlations
        st.subheader("🎯 Brentwood Fill Selection")

        # Restrict fill options based on tower arrangement (very important for correlation validity)
        if tower_type == "crossflow":
            fill_options = ["XF75", "CF1900SBMA", "XF125", "XF3000"]
            default_fills = ["XF75", "CF1900SBMA"]
        else:
            # counterflow_induced / counterflow_forced
            fill_options = ["CT1200AT", "CF1200", "CF1900SBMA", "XF3000"]
            default_fills = ["CT1200AT", "CF1900SBMA", "CF1200"]

        # Keep only fills that exist in the database (defensive)
        fill_options = [f for f in fill_options if f in BRENTWOOD_FILLS]
        default_fills = [f for f in default_fills if f in fill_options]

        selected_fills = st.multiselect(
            "Select fills to compare:",
            options=fill_options,
            default=default_fills,
            format_func=lambda x: BRENTWOOD_FILLS[x]["name"],
            help="Select one or more fills for comparison"
        )

        # NEW: PDF Data Usage Option for CT1200AT
        if "CT1200AT" in selected_fills and tower_type.startswith("counterflow"):
            st.subheader("📊 CT1200AT Data Source")
            use_pdf_data = st.radio(
                "CT1200AT Performance Data Source:",
                ["Use PDF Graph Data (2017)", "Use Default Performance Data"],
                help="PDF data includes specific Ka/L and pressure drop curves from 2017 document"
            )
            use_pdf_data_bool = (use_pdf_data == "Use PDF Graph Data (2017)")
            st.session_state.use_pdf_data = use_pdf_data_bool
            
            if use_pdf_data_bool:
                st.info("""
                **Using Brentwood Published Correlations (CT1200AT, 2017):**
                - Heat transfer uses Brentwood's *KaV/L* equation from the performance sheet
                - Pressure drop uses Brentwood's published ΔP correlation (with unit conversions)
                - **Counterflow-only:** these CT1200AT correlations are validated for counterflow arrangement

                **Validity guidance (avoid extrapolation):**
                - Fill height (tested): **0.61–1.22 m** (610–1220 mm)
                - L/G ratio (plotted): **0.5–2.5**
                - Water loading (plotted): **0–30 m³/h·m²**
                - Air face velocity (plotted): **~1–4 m/s**

                The app will warn if you drift outside these ranges.
                """)
            else:
                st.info("""
                **Using Default Performance Data:**
                - Estimated performance similar to CF1200 but slightly better
                - Good for preliminary design comparisons
                """)
        
        # NEW: Supplier Validation Button
        
        # NEW: PDF Data Usage Option for XF75 (Crossflow only)
        if "XF75" in selected_fills and tower_type == "crossflow":
            st.subheader("📊 XF75 Data Source")
            use_xf75_pdf = st.radio(
                "XF75 Performance Data Source:",
                ["Use PDF Graph Data (2018)", "Use Default Performance Data"],
                help="PDF data includes published KaH/L and pressure drop correlations from 2018 document"
            )
            use_xf75_pdf_bool = (use_xf75_pdf == "Use PDF Graph Data (2018)")
            st.session_state.use_xf75_pdf_data = use_xf75_pdf_bool

            h_sel, at_sel, w_sel, bank_sel = xf75_crossflow_geometry_from_session()
            st.caption(f"Selected XF75 geometry: H = {h_sel:.3f} m, AT = {at_sel:.3f} m, stack width = {w_sel:.3f} m")

            if use_xf75_pdf_bool:
                st.info("""
                **Using Brentwood Published Correlations (XF75, 2018):**
                - Heat transfer uses Brentwood's *KaH/L* equation (SI form)
                - **UI input remains L/G.** Brentwood SI equation uses G/L, so the app internally computes **G/L = 1/(L/G)**
                - Pressure drop uses Brentwood's published ΔP correlation (SI form)
                - **Crossflow:** XF75 is a crossflow herringbone film fill (do not use in counterflow)

                **Validity guidance (avoid extrapolation):**
                - Fill height H on Brentwood plots: 1.829–4.877 m (6–16 ft)
                - Air travel depth AT on plots: 0.610–1.829 m (2–6 ft)
                - Air velocity on ΔP plots: ~1.0–4.0 m/s
                - Water loading QA on ΔP plots: ~20–70 m³/(h·m²)
                - L/G on Ka plots: roughly 0.4–3.0
                """)
        st.subheader("🔍 Supplier Validation")
        run_saa15_validation = st.button(
            "🔄 Run SAA15 Supplier Design Validation",
            help="Compare your code against supplier's SAA15 CF1200 design",
            use_container_width=True
        )
        
        # Run Button
        run_calc = st.button("🚀 Run Complete Analysis", type="primary", use_container_width=True)
        
        # Report generation
        st.subheader("📄 Report Generation")
        generate_reports = st.checkbox("Generate TXT Report", value=True)
    
    # ========================================================================
    # MAIN CONTENT
    # ========================================================================
    
    # NEW: SAA15 Supplier Validation Section
    if run_saa15_validation:
        st.header("🔬 SAA15 Supplier Design Validation")
        st.info("""
        **Validating against supplier's SAA15 design with CF1200 fill:**
        - Water flow: 114 kg/s
        - L/G ratio: 2.313
        - Fill: CF1200, depth 0.75m
        - Tower: Counterflow induced draft
        - Hot water: 40°C, Cold target: 35°C
        - Wet bulb: 30°C, Dry bulb: 33°C
        """)
        
        with st.spinner("Running validation against supplier's SAA15 design..."):
            results, comparison = validate_with_saa15_supplier_design()
        
        # Display comparison
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Your Code's Calculation")
            st.metric("Cold Water Temp", f"{results['T_cold_achieved']:.2f}°C")
            st.metric("Fan Power", f"{results['fan_power']:.2f} kW")
            st.metric("Ka/L", f"{results['Ka_over_L']:.3f}")
            st.metric("K Value (Ka)", f"{results['K_value']:.2f} kW/°C")  # Added K value
            st.metric("Static Pressure", f"{results['total_static_pressure']:.1f} Pa")
            st.metric("Water Loading", f"{results['water_loading']:.1f} m³/h·m²")
            st.metric("Air Density", f"{results['air_density']:.3f} kg/m³")
        
        with col2:
            st.subheader("Supplier's Claim (SAA15)")
            st.metric("Cold Water Temp", "35.00°C")
            st.metric("Fan Power", "13.41 kW")
            st.metric("Ka/L", "0.982")
            st.metric("K Value (Ka)", "112.0 kW/°C")  # Calculated: 0.982 * 114
            st.metric("Static Pressure", f"{comparison['supplier_claimed']['static_pressure_Pa']:.1f} Pa")
            st.metric("Water Loading", f"{comparison['supplier_claimed']['water_loading']:.1f} m³/h·m²")
            st.metric("Air Density", "~0.915 kg/m³")
        
        # Differences
        st.subheader("📊 Differences")
        diff_col1, diff_col2, diff_col3, diff_col4, diff_col5 = st.columns(5)
        with diff_col1:
            delta_temp = comparison['differences']['T_cold_diff']
            st.metric("Δ Cold Temp", f"{delta_temp:.2f}°C", 
                     delta_color="inverse" if delta_temp > 0 else "normal")
        with diff_col2:
            delta_power = comparison['differences']['fan_power_diff']
            st.metric("Δ Fan Power", f"{delta_power:.2f} kW",
                     delta_color="inverse" if delta_power > 0 else "normal")
        with diff_col3:
            delta_kal = comparison['differences']['Ka_over_L_diff']
            st.metric("Δ Ka/L", f"{delta_kal:.3f}",
                     delta_color="normal" if delta_kal > 0 else "inverse")
        with diff_col4:
            k_value = results['K_value']
            expected_k = 0.982 * 114
            delta_k = k_value - expected_k
            st.metric("Δ K Value", f"{delta_k:.1f} kW/°C",
                     delta_color="normal" if delta_k > 0 else "inverse")
        with diff_col5:
            st.metric("RH Calculated", f"{results['RH']:.1f}%", "From dry/wet bulb")
        
        # Interpretation
        st.info("""
        **Interpretation:**
        - If your code matches supplier closely (±5%), CF1200 curve is accurate
        - If your code shows better performance, supplier may be conservative
        - If your code shows worse performance, check pressure drop assumptions
        - Dry bulb temperature affects air density and psychrometric calculations
        """)
        
        # Show detailed results
        with st.expander("📋 Detailed Validation Results"):
            st.write("**Your Calculation:**")
            st.json({k: round(v, 3) if isinstance(v, float) else v 
                    for k, v in comparison['your_calculation'].items()})
    
    # ========================================================================
    # MAIN CALCULATION
    # ========================================================================
    if run_calc and selected_fills:
        # Validate temperatures
        if T_hot <= T_cold_target:
            st.error("❌ Error: Hot water temperature must be GREATER than cold water target")
            st.stop()
        
        # Validate dry bulb is greater than wet bulb
        if Tdb <= Twb:
            st.warning("⚠️ Note: Dry bulb temperature should typically be higher than wet bulb temperature")
        
        # Calculate for all selected fills
        results = []
        
        with st.spinner("Running cooling tower calculations..."):

            # Brentwood correlation range guidance (CT1200AT, counterflow)
            if "CT1200AT" in selected_fills and st.session_state.get("use_pdf_data", False):
                # Compute key operating points (same definitions used inside the solver)
                try:
                    air_density_tmp = air_density_calc(Tdb, Twb, altitude)
                    air_flow_vol_tmp = G / air_density_tmp  # m³/s
                    air_face_velocity_tmp = air_flow_vol_tmp / st.session_state.face_area  # m/s
                    water_loading_tmp = (L * 3.6) / st.session_state.face_area  # m³/h·m² (assumes ρw≈1000)
                    L_over_G_tmp = L / G if G != 0 else float("inf")
                except Exception:
                    air_face_velocity_tmp = None
                    water_loading_tmp = None
                    L_over_G_tmp = None

                if not tower_type.startswith("counterflow"):
                    st.warning("⚠️ CT1200AT Brentwood correlations are validated for **counterflow** towers. "
                               "You selected a non-counterflow tower type; the app will fall back to the digitized curve method for ΔP "
                               "and will not apply the KaV/L equation.")

                # Range checks (warn only; do not stop the run)
                def _warn_outside(name, val, vmin, vmax, unit=""):
                    if val is None or (isinstance(val, float) and (math.isnan(val) or math.isinf(val))):
                        return
                    if val < vmin or val > vmax:
                        st.warning(f"⚠️ CT1200AT correlation: {name} = {val:.3g}{unit} is outside the plotted/tested range "
                                   f"({vmin:g}–{vmax:g}{unit}). Extrapolation may be unreliable.")

                _warn_outside("Fill height", fill_depth, 0.61, 1.22, " m")
                _warn_outside("L/G ratio", L_over_G_tmp, 0.5, 2.5, "")
                _warn_outside("Water loading", water_loading_tmp, 0.0, 30.0, " m³/h·m²")
                _warn_outside("Air face velocity", air_face_velocity_tmp, 1.0, 4.0, " m/s")

            # Brentwood XF75 crossflow SI chart range guidance
            if "XF75" in selected_fills and tower_type == "crossflow" and st.session_state.get("use_xf75_pdf_data", False):
                try:
                    h_xf75, at_xf75, width_xf75, bank_count_xf75 = xf75_crossflow_geometry_from_session()
                    air_density_tmp = air_density_calc(Tdb, Twb, altitude)
                    air_flow_vol_tmp = G / air_density_tmp
                    air_face_area_tmp = bank_count_xf75 * h_xf75 * width_xf75
                    water_plan_area_tmp = bank_count_xf75 * at_xf75 * width_xf75
                    air_face_velocity_tmp = air_flow_vol_tmp / air_face_area_tmp
                    water_loading_tmp = (L * 3.6) / water_plan_area_tmp
                    L_over_G_tmp = L / G if G != 0 else float("inf")
                except Exception:
                    h_xf75 = at_xf75 = air_face_velocity_tmp = water_loading_tmp = L_over_G_tmp = None

                def _warn_xf75(name, val, vmin, vmax, unit=""):
                    if val is None or (isinstance(val, float) and (math.isnan(val) or math.isinf(val))):
                        return
                    if val < vmin or val > vmax:
                        st.warning(f"⚠️ XF75 SI charts: {name} = {val:.3g}{unit} is outside the published graph range "
                                   f"({vmin:g}–{vmax:g}{unit}). Extrapolation may be unreliable.")

                _warn_xf75("Fill height H", h_xf75, min(XF75_SI_FILL_HEIGHT_OPTIONS_M), max(XF75_SI_FILL_HEIGHT_OPTIONS_M), " m")
                _warn_xf75("Air travel depth AT", at_xf75, min(XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M), max(XF75_SI_AIR_TRAVEL_DEPTH_OPTIONS_M), " m")
                _warn_xf75("L/G ratio", L_over_G_tmp, XF75_SI_LG_RANGE[0], XF75_SI_LG_RANGE[1], "")
                _warn_xf75("Water loading QA", water_loading_tmp, XF75_SI_WATER_LOADING_RANGE[0], XF75_SI_WATER_LOADING_RANGE[1], " m³/h·m²")
                _warn_xf75("Air face velocity", air_face_velocity_tmp, XF75_SI_AIR_VELOCITY_RANGE[0], XF75_SI_AIR_VELOCITY_RANGE[1], " m/s")

            for fill in selected_fills:
                use_pdf = False
                if fill == "CT1200AT":
                    use_pdf = bool(st.session_state.get("use_pdf_data", False))
                elif fill == "XF75":
                    use_pdf = bool(st.session_state.get("use_xf75_pdf_data", False))
                
                result = solve_cooling_tower_enhanced(
                    L, G, T_hot, T_cold_target, Twb, Tdb, fill,
                    tower_type, fill_depth, st.session_state.face_area, altitude, use_pdf
                )
                results.append(result)
        
        # Display results
        st.header("📊 Performance Results")
        
        # Create metrics columns
        cols = st.columns(len(selected_fills))
        for idx, (col, result) in enumerate(zip(cols, results)):
            with col:
                st.subheader(f"{result['fill_name']}")
                st.caption(f"{result['tower_name']}")
                if result.get('use_pdf_data', False):
                    st.caption("📊 Using PDF Data")
                
                temp_status = "✅" if result['T_cold_achieved'] <= result['T_cold_target'] else "❌"
                st.metric(f"{temp_status} Cold Water", 
                         f"{result['T_cold_achieved']:.2f}°C",
                         delta=f"{result['T_cold_achieved'] - result['T_cold_target']:.2f}°C vs target")
                
                st.metric("Heat Rejection", f"{result['Q_achieved']:.0f} kW")
                st.metric("Fan Power", f"{result['fan_power']:.2f} kW")
                st.metric("K Value (Ka)", f"{result['K_value']:.1f} kW/°C")  # Added K value
                st.metric("Static Pressure", f"{result['total_static_pressure']:.0f} Pa")
                st.metric("Water Loading", f"{result['water_loading']:.1f} m³/h·m²")
        
        # Display atmospheric conditions
        st.subheader("🌤️ Atmospheric Conditions")
        col_atm1, col_atm2, col_atm3, col_atm4 = st.columns(4)
        with col_atm1:
            st.metric("Dry Bulb", f"{Tdb:.1f}°C")
        with col_atm2:
            st.metric("Wet Bulb", f"{Twb:.1f}°C")
        with col_atm3:
            # Use first result for RH (same for all fills)
            st.metric("Relative Humidity", f"{results[0]['RH']:.1f}%")
        with col_atm4:
            st.metric("Altitude", f"{altitude} m ASL")
        
        # Brentwood XF75 graph validity display
        xf75_checked_results = [r for r in results if r.get('xf75_graph_checks')]
        if xf75_checked_results:
            st.subheader("🧭 Brentwood XF75 SI Graph Validity")
            for r in xf75_checked_results:
                checks = r['xf75_graph_checks']
                if checks.get('all_ok'):
                    st.success(f"{r['fill_name']}: all checked points are inside the Brentwood XF75 SI graph envelopes.")
                else:
                    st.warning(f"{r['fill_name']}: one or more points are outside the Brentwood XF75 SI graph envelopes; results may be extrapolated.")
                with st.expander(f"Show XF75 graph checks — {r['fill_name']}"):
                    for label, ok, message in checks.get('checks', []):
                        if ok:
                            st.write(f"✅ **{label}** — {message}")
                        else:
                            st.write(f"⚠️ **{label}** — {message}")

        # Detailed Comparison Table
        st.header("📋 Detailed Performance Comparison")
        comparison_data = []
        for result in results:
            comparison_data.append({
                "Fill Type": result['fill_name'],
                "Tower Type": result['tower_name'],
                "Data Source": "PDF" if result.get('use_pdf_data', False) else "Default",
                "Cold Water (°C)": f"{result['T_cold_achieved']:.2f}",
                "Required Heat Rejection (kW)": f"{result.get('Q_target', 0):.0f}",
                "Achievable Heat Rejection (kW)": f"{result['Q_achieved']:.0f}",
                "Heat Margin (kW)": f"{result['Q_achieved'] - result.get('Q_target', 0):.0f}",
                "K Value (kW/°C)": f"{result['K_value']:.1f}",  # Added K value
                "Ka/L (1/m)": f"{result['Ka_over_L']:.3f}",
                "KaH/L or KaV/L": f"{result.get('KaH_over_L', result['Ka_over_L'] * result['fill_depth']):.3f}",
                "Approach (°C)": f"{result['approach']:.2f}",
                "Range (°C)": f"{result['cooling_range']:.2f}",
                "L/G Ratio": f"{result['L_over_G']:.3f}",
                "NTU": f"{result['NTU']:.3f}",
                "Surface Area (m²)": f"{result['total_surface_area']:.0f}",
                "Surface Area Density (m²/m³)": f"{result['surface_area_density']:.1f}",
                "Fill Banks": f"{result.get('crossflow_geometry', {}).get('crossflow_bank_count', 1) if result.get('tower_type') == 'crossflow' else '-'}",
                "Air Face Area (m²)": f"{result['face_area']:.2f}",
                "Net Open Air Area (m²)": f"{result.get('open_air_area', 0):.2f}",
                "Water Plan Area (m²)": f"{result.get('water_plan_area', result['face_area']):.2f}",
                "Water Velocity in Channels (m/s)": f"{result['water_velocity']:.3f}",
                "Film Thickness (mm)": f"{result['water_film_thickness']}",
                "Water Loading (m³/h·m²)": f"{result['water_loading']:.1f}",
                "Air Face Velocity (m/s)": f"{result['air_face_velocity']:.2f}",
                "Air Velocity Through Fill (m/s)": f"{result.get('air_velocity_through_fill', result['air_face_velocity']):.2f}",
                "Free Area Fraction": f"{result.get('free_area_fraction', 0):.2f}",
                "XF75 Graph Range": ("OK" if (not result.get('xf75_graph_checks') or result['xf75_graph_checks'].get('all_ok')) else "CHECK WARNING"),
                "Air Density (kg/m³)": f"{result['air_density']:.3f}",
                "Fill ΔP (Pa)": f"{result['fill_pressure_drop']:.1f}",
                "Other Losses (Pa)": f"{result.get('other_losses_total', 0):.1f}",
                "Fan Power (kW)": f"{result['fan_power']:.2f}",
                "Static Pressure (Pa)": f"{result['total_static_pressure']:.0f}",
                "Fouling Risk": result['fouling_risk']['risk_level']
            })
        
        df_comparison = pd.DataFrame(comparison_data)
        st.dataframe(df_comparison, use_container_width=True)

        # Download comparison table as CSV (explicit button; avoids browser/table widget quirks)
        csv_bytes = df_comparison.to_csv(index=False).encode("utf-8")
        st.download_button(
            label="📥 Download Comparison Table (CSV)",
            data=csv_bytes,
            file_name=f"cooling_tower_comparison_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv",
            use_container_width=True
        )
        
        # Tower Geometry Summary
        st.header("📐 Tower Geometry Summary")
        if tower_type == "crossflow":
            h_cf, at_cf, w_cf, bank_count_cf = xf75_crossflow_geometry_from_session()
            air_area_single = h_cf * w_cf
            water_area_single = at_cf * w_cf
            fill_volume_single = h_cf * at_cf * w_cf
            air_area_summary = bank_count_cf * air_area_single
            water_area_summary = bank_count_cf * water_area_single
            fill_volume_summary = bank_count_cf * fill_volume_single
            col_geo1, col_geo2, col_geo3, col_geo4 = st.columns(4)
            with col_geo1:
                st.metric("Fill Banks", f"{bank_count_cf}")
            with col_geo2:
                st.metric("Total Air Face Area", f"{air_area_summary:.2f} m²")
            with col_geo3:
                st.metric("Total Water Plan Area", f"{water_area_summary:.2f} m²")
            with col_geo4:
                st.metric("Total Fill Volume", f"{fill_volume_summary:.2f} m³")
            st.caption(
                f"Per bank geometry: H × AT × Width = {h_cf:.2f} × {at_cf:.2f} × {w_cf:.2f} m; "
                f"per-bank air face area = {air_area_single:.2f} m²."
            )
            if any(r.get('fill_type') == 'XF75' for r in results):
                st.info(
                    f"XF75 catalogue interpretation: D=AT {at_cf:.3f} m, W=stack width per bank {w_cf:.3f} m, "
                    f"L=vertical fill height {h_cf:.3f} m, banks={bank_count_cf}. Estimated total wetted surface area at 167.4 m²/m³ = "
                    f"{fill_volume_summary * XF75_CATALOG_SURFACE_AREA_M2_M3:,.0f} m²."
                )
        else:
            col_geo1, col_geo2, col_geo3, col_geo4 = st.columns(4)
            geom = st.session_state.get("geometry_input", {})
            with col_geo1:
                st.metric("Face Area", f"{st.session_state.face_area:.2f} m²")
            with col_geo2:
                st.metric("Fill Depth", f"{fill_depth:.3f} m")
            with col_geo3:
                st.metric("Tower Shape", tower_shape)
            with col_geo4:
                st.metric("Fill Volume", f"{st.session_state.face_area * fill_depth:.2f} m³")
            if geom.get("tower_shape") == "Rectangle":
                st.caption(f"Input rectangular fill face: Length = {geom.get('fill_length_m', 0):.3f} m, Width/Breadth = {geom.get('fill_width_m', 0):.3f} m.")
            elif geom.get("tower_shape") == "Round":
                st.caption(f"Input round tower diameter = {geom.get('diameter_m', 0):.3f} m.")

        current_effective_kavl, _ = back_calculate_effective_kavl(T_hot, T_cold_target, Twb, tower_type)
        render_nearest_specs_panel({
            'heat_load_kw': Q_input,
            'water_flow_m3h': water_flow_m3h,
            'hot_water_c': T_hot,
            'cold_water_c': T_cold_target,
            'wet_bulb_c': Twb,
            'length_m': fill_length if tower_shape == "Rectangle" else None,
            'width_m': fill_width if tower_shape == "Rectangle" else None,
            'effective_kavl': current_effective_kavl
        })
        
        # Tower Type Analysis
        st.header("🏗️ Tower Type Analysis")
        tower_desc = TOWER_TYPES[tower_type]
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Fill Utilization", f"{tower_desc['fill_utilization']*100:.0f}%")
        with col2:
            st.metric("Pressure Drop Factor", f"{tower_desc['typical_pressure_drop_factor']:.1f}x")
        with col3:
            st.metric("Air-Water Contact", tower_desc['air_water_contact'])
        
        st.info(f"**{tower_desc['description']}**")
        
        # Visualization
        st.header("📈 Performance Visualization")
        
        fig, (ax1, ax2, ax3) = plt.subplots(1, 3, figsize=(15, 5))
        
        # Cold water temperatures
        fill_names = [r['fill_name'][:15] + (" (PDF)" if r.get('use_pdf_data', False) else "") for r in results]
        cold_temps = [r['T_cold_achieved'] for r in results]
        
        bars1 = ax1.bar(fill_names, cold_temps, color=['red' if t > T_cold_target else 'green' for t in cold_temps])
        ax1.axhline(y=T_cold_target, color='blue', linestyle='--', label='Target')
        ax1.set_ylabel('Cold Water Temperature (°C)')
        ax1.set_title('Performance vs Target')
        ax1.tick_params(axis='x', rotation=45)
        ax1.legend()
        
        # Fan power comparison
        fan_powers = [r['fan_power'] for r in results]
        bars2 = ax2.bar(fill_names, fan_powers, color='orange', alpha=0.7)
        ax2.set_ylabel('Fan Power (kW)')
        ax2.set_title('Energy Consumption')
        ax2.tick_params(axis='x', rotation=45)
        
        # K value comparison
        k_values = [r['K_value'] for r in results]
        bars3 = ax3.bar(fill_names, k_values, color='purple', alpha=0.7)
        ax3.set_ylabel('K Value (kW/°C)')
        ax3.set_title('Heat Transfer Coefficient (K)')
        ax3.tick_params(axis='x', rotation=45)
        
        # Add value labels
        for bars, ax in [(bars1, ax1), (bars2, ax2), (bars3, ax3)]:
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + (0.02 * max([b.get_height() for b in bars])),
                       f'{height:.1f}' if ax == ax1 else f'{height:.1f}',
                       ha='center', va='bottom', fontsize=8)
        
        st.pyplot(fig)
        
        # CT1200AT Specific Notes
        if "CT1200AT" in selected_fills and tower_type.startswith("counterflow"):
            st.header("📝 CT1200AT Performance Data Information")
            col_info1, col_info2 = st.columns(2)
            
            with col_info1:
                st.subheader("📊 PDF Graph Data (2017)")
                st.markdown("""
                **Performance Characteristics:**
                - Surface Area: 226 m²/m³
                - Hydraulic Diameter: 8.8 mm
                - Flute Angle: 30°
                - Free Area: 89%
                
                **Pressure Drop (varies by fill height):**
                - 24" (610mm): Base curve for shallow fill
                - 36" (915mm): Medium depth
                - 48" (1220mm): Deep fill applications
                
                **Ka/L Curve:**
                - Based on interpolation of PDF performance graphs
                - Slightly better than CF1200
                - More accurate than default data
                """)
            
            with col_info2:
                st.subheader("⚖️ Comparison with CF1200")
                st.markdown("""
                **CT1200AT vs CF1200:**
                - **Ka/L Values:** CT1200AT ~10% higher
                - **Pressure Drop:** CT1200AT slightly lower
                - **Fouling Resistance:** CT1200AT better (85% vs 80%)
                - **Data Source:** Actual PDF graphs vs estimated
                
                **When to use PDF data:**
                - For accurate design verification
                - When comparing with existing installations
                - For final design calculations
                
                **When to use default data:**
                - For preliminary sizing
                - When exact PDF curves are not critical
                - For quick comparisons
                """)
        
        # CF1200 Specific Notes
        if "CF1200" in selected_fills:
            st.warning("""
            **Note about CF1200 fill:**
            - This is an **older fill design** with performance data tuned to match supplier's SAA15 design
            - Compared to modern fills (XF75), CF1200 has:
              - Lower thermal efficiency (Ka/L ~70% of XF75)
              - Higher pressure drop
              - Worse fouling resistance
            - Use for comparison with existing towers or supplier designs
            """)
        
        # Report Generation
        if generate_reports:
            st.header("📄 Report Generation")
            selected_for_report = st.selectbox(
                "Select design for detailed report:",
                options=[f"{r['fill_name']} ({r['T_cold_achieved']:.2f}°C)" for r in results],
                index=0
            )
            
            fill_index = [f"{r['fill_name']} ({r['T_cold_achieved']:.2f}°C)" for r in results].index(selected_for_report)
            selected_result = results[fill_index]
            
            txt_report = generate_txt_report(selected_result)
            st.download_button(
                label="📥 Download Detailed TXT Report",
                data=txt_report,
                file_name=f"cooling_tower_{selected_result['fill_type']}_{selected_result['tower_type']}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                mime="text/plain",
                use_container_width=True
            )

            # A4 PDF report for easy printing
            pdf_bytes = generate_pdf_report_from_text(
                txt_report,
                title=f"Cooling Tower Report - {selected_result['fill_name']} ({selected_result['tower_name']})"
            )
            st.download_button(
                label="📥 Download A4 PDF Report",
                data=pdf_bytes,
                file_name=f"cooling_tower_{selected_result['fill_type']}_{selected_result['tower_type']}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
            
            # Show report preview
            with st.expander("📋 Preview Report (First 2000 characters)"):
                st.text(txt_report[:2000] + "..." if len(txt_report) > 2000 else txt_report)
    
    elif run_calc and not selected_fills:
        st.warning("Please select at least one Brentwood fill type.")
    else:
        # Welcome message
        st.markdown("""
        ## 🌊 Enhanced Cooling Tower Design Tool
        
        ### ✅ **NEW FEATURES ADDED:**
        
        1. **CT1200AT Fill Added:**
           - New fill type from Brentwood PDF data
           - Actual performance data from 2017 graphs
           - Radio button to choose data source (PDF vs Default)
           - Pressure drop varies by fill height (24", 36", 48")
        
        2. **K Value Display:**
           - **K value (Ka)** now shown in all results
           - Displayed as "K Value (kW/°C)" 
           - Added to comparison tables and reports
           - Included in validation section
        
        3. **Enhanced Comparison:**
           - Clear distinction between PDF and default data
           - Better visualization with K value graphs
           - More detailed fill information
        
        ### 📊 **K Value Information:**
        
        The **K value (Ka)** represents the **overall heat transfer coefficient** in kW/°C:
        - **Formula:** K = Ka = (Ka/L) × Water Flow Rate (L)
        - **Physical Meaning:** Heat transfer per degree temperature difference
        - **Units:** kW/°C
        - **Higher K value** = Better heat transfer capability
        
        ### 🎯 **How to Use the New Features:**
        
        1. **Select CT1200AT** from fill options
        2. **Choose data source:** PDF Graph Data or Default Data
        3. **Run analysis** to see K values for all fills
        4. **Compare CT1200AT** with other fills
        5. **Check validation** against supplier SAA15 design
        
        ### 🔍 **CT1200AT Data Sources:**
        
        | Feature | PDF Graph Data | Default Data |
        |---------|----------------|--------------|
        | Ka/L Values | From actual graphs | Estimated |
        | Pressure Drop | Based on fill height | Standard curve |
        | Accuracy | High (actual data) | Medium (estimated) |
        | Use Case | Final design | Preliminary sizing |
        
        ---
        
        *Configure your design in the sidebar using the enhanced controls.*
        """)



# ============================================================================
# GOOGLE DRIVE BACKEND (STREAMLIT CLOUD FRIENDLY)
# ============================================================================

def google_drive_is_configured():
    try:
        return ('gcp_service_account' in st.secrets) and ('google_drive_folder_id' in st.secrets)
    except Exception:
        return False


def get_drive_service():
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    creds_info = dict(st.secrets['gcp_service_account'])
    creds = service_account.Credentials.from_service_account_info(
        creds_info,
        scopes=['https://www.googleapis.com/auth/drive.file']
    )
    return build('drive', 'v3', credentials=creds)


def _drive_find_file(service, filename):
    folder_id = st.secrets['google_drive_folder_id']
    safe_name = filename.replace("'", "\\'")
    q = f"name = '{safe_name}' and '{folder_id}' in parents and trashed = false"
    resp = service.files().list(q=q, fields='files(id,name,modifiedTime)', pageSize=20).execute()
    files = resp.get('files', [])
    if not files:
        return None
    files = sorted(files, key=lambda x: x.get('modifiedTime', ''), reverse=True)
    return files[0]


def upload_bytes_to_drive(filename: str, data: bytes, mime_type: str):
    from googleapiclient.http import MediaIoBaseUpload
    service = get_drive_service()
    folder_id = st.secrets['google_drive_folder_id']
    existing = _drive_find_file(service, filename)
    media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mime_type, resumable=False)
    if existing:
        updated = service.files().update(fileId=existing['id'], media_body=media, fields='id,name,webViewLink').execute()
        return updated
    body = {'name': filename, 'parents': [folder_id]}
    created = service.files().create(body=body, media_body=media, fields='id,name,webViewLink').execute()
    return created


def read_drive_file_bytes(filename: str):
    from googleapiclient.http import MediaIoBaseDownload
    service = get_drive_service()
    meta = _drive_find_file(service, filename)
    if not meta:
        return None
    request = service.files().get_media(fileId=meta['id'])
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()
    fh.seek(0)
    return fh.read()


def upload_json_to_drive(filename: str, data_dict: dict):
    payload = json.dumps(data_dict, indent=2, ensure_ascii=False).encode('utf-8')
    return upload_bytes_to_drive(filename, payload, 'application/json')


def upload_text_to_drive(filename: str, text: str):
    return upload_bytes_to_drive(filename, text.encode('utf-8'), 'text/plain')


def _read_library():
    if google_drive_is_configured():
        try:
            raw = read_drive_file_bytes('vendor_specs_library.csv')
            if raw:
                return pd.read_csv(io.BytesIO(raw))
        except Exception:
            pass
    if LIBRARY_PATH.exists():
        try:
            return pd.read_csv(LIBRARY_PATH)
        except Exception:
            return pd.DataFrame()
    return pd.DataFrame()


def _write_library(df):
    wrote = False
    if google_drive_is_configured():
        try:
            csv_bytes = df.to_csv(index=False).encode('utf-8')
            upload_bytes_to_drive('vendor_specs_library.csv', csv_bytes, 'text/csv')
            wrote = True
        except Exception:
            wrote = False
    try:
        df.to_csv(LIBRARY_PATH, index=False)
    except Exception:
        pass
    return wrote or LIBRARY_PATH.exists()


def save_spec_to_library(spec_dict):
    row = pd.DataFrame([spec_dict])
    existing = _read_library()
    combined = pd.concat([existing, row], ignore_index=True) if not existing.empty else row
    return _write_library(combined)


def save_uploaded_files_to_drive(uploaded_files):
    saved = []
    for uf in uploaded_files or []:
        try:
            data = uf.getvalue()
            result = upload_bytes_to_drive(uf.name, data, uf.type or 'application/octet-stream')
            saved.append(result)
        except Exception:
            continue
    return saved


def cache_uploaded_vendor_files(uploaded_files):
    """Store uploaded vendor files safely in Streamlit session state so button clicks/reruns don't lose them."""
    cached = []
    for uf in uploaded_files or []:
        try:
            cached.append({
                'name': uf.name,
                'type': uf.type or 'application/octet-stream',
                'bytes': uf.getvalue(),
            })
        except Exception:
            continue
    st.session_state['vendor_uploaded_files_cached'] = cached
    return cached


def get_cached_vendor_files():
    return st.session_state.get('vendor_uploaded_files_cached', [])


def save_cached_vendor_files_to_drive():
    saved = []
    errors = []
    for item in get_cached_vendor_files():
        try:
            result = upload_bytes_to_drive(item['name'], item['bytes'], item.get('type') or 'application/octet-stream')
            saved.append(result)
        except Exception as e:
            errors.append(f"{item.get('name','unknown')}: {e}")
    return saved, errors


def render_vendor_comparison_mode():
    st.header('📄 Vendor Quote Analysis & Automatic Fill Comparison')
    st.caption('Upload up to 5 vendor PDF/DOCX/TXT files. The app combines them, extracts technical data, converts to SI, compares Brentwood fills, and can save the captured spec to a reusable library.')
    if 'vendor_uploaded_files_cached' not in st.session_state:
        st.session_state['vendor_uploaded_files_cached'] = []
    if google_drive_is_configured():
        st.success('Google Drive storage is configured. Parsed specs library and uploaded vendor files can be saved to your shared Drive folder.')
    else:
        st.info('Google Drive storage is not configured yet. The app will fall back to local CSV storage where possible. On Streamlit Cloud, local files may not persist across restarts.')

    with st.sidebar:
        st.header('📎 Vendor Comparison Inputs')
        uploaded_files = st.file_uploader('Upload vendor documents', type=['pdf','docx','txt'], accept_multiple_files=True, key='vendor_docs_uploader')
        default_tower_type = st.selectbox('Default tower type if parser is unsure', options=list(TOWER_TYPES.keys()), format_func=lambda x: TOWER_TYPES[x]['name'], index=1)
        default_fill_depth = st.number_input('Assumed fill depth for comparison (m)', min_value=0.300, max_value=2.500, value=1.000, step=0.050, format='%.3f')
        default_altitude = st.number_input('Site altitude (m)', min_value=0, max_value=5000, value=0, step=50)
        use_pdf_for_ct1200 = st.checkbox('Use Brentwood CT1200AT published equations', value=True)
        use_pdf_for_xf75 = st.checkbox('Use Brentwood XF75 published equations', value=True)
        auto_save_to_library = st.checkbox('Auto-save parsed spec to library after comparison', value=google_drive_is_configured())
        auto_save_originals = st.checkbox('Auto-save uploaded vendor files to Google Drive', value=False, disabled=not google_drive_is_configured())
        test_drive = st.button('☁️ Test Google Drive connection', use_container_width=True, disabled=not google_drive_is_configured())
        run_compare = st.button('🚀 Run Vendor Comparison', type='primary', use_container_width=True)

    if test_drive and google_drive_is_configured():
        try:
            ts = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            r = upload_text_to_drive('drive_test.txt', f'Cooling tower Drive test successful. {ts}')
            st.success(f"Drive upload successful: {r.get('name')}")
            if r.get('webViewLink'):
                st.write(r.get('webViewLink'))
        except Exception as e:
            st.error(f'Drive upload failed: {e}')

    extracted = {}
    combined_text = ''
    if uploaded_files:
        if len(uploaded_files) > 5:
            st.error('Please upload a maximum of 5 files at one time.')
            return
        cache_uploaded_vendor_files(uploaded_files)
        combined_text = extract_text_from_multiple_files(uploaded_files)
        extracted = parse_vendor_quote_text(combined_text)
        st.subheader('🔎 Extracted Technical Data (best effort)')
        preview_rows = []
        for k, v in extracted.items():
            if k == 'source_summary':
                continue
            preview_rows.append({'Parameter': k, 'Value': v})
        preview = pd.DataFrame(preview_rows)
        st.dataframe(preview, use_container_width=True, height=350)
        if extracted.get('source_summary'):
            with st.expander('Show extraction notes'):
                for item in extracted['source_summary']:
                    st.write('-', item)
        cap_docx = generate_docx_report('Captured Vendor Technical Data', [('Uploaded Files', [uf.name for uf in uploaded_files]), ('Extracted Data', {k:v for k,v in extracted.items() if k!='source_summary'}), ('Extraction Notes', extracted.get('source_summary', []))])
        st.download_button('📥 Download Captured Data as Word (.docx)', data=cap_docx, file_name=f'captured_vendor_data_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', use_container_width=True)
        with st.expander('🛠️ Upload/Drive debug'):
            st.write('Google Drive configured:', google_drive_is_configured())
            st.write('Uploader returned files:', len(uploaded_files or []))
            st.write('Cached files in session:', len(get_cached_vendor_files()))
            st.write('Cached names:', [x.get('name') for x in get_cached_vendor_files()])
        if google_drive_is_configured() and st.button('☁️ Save uploaded vendor files to Google Drive', use_container_width=True):
            cached_vendor_files = get_cached_vendor_files()
            if not cached_vendor_files:
                st.warning('No uploaded vendor files are available in session cache. Please upload the files again and then click save.')
            else:
                try:
                    saved, errors = save_cached_vendor_files_to_drive()
                    if saved:
                        st.success(f"Saved {len(saved)} uploaded file(s) to Google Drive: " + ', '.join([x.get('name','?') for x in saved]))
                    else:
                        st.warning('Drive save attempted, but no files were saved.')
                    if errors:
                        st.error('Some files failed: ' + ' | '.join(errors))
                except Exception as e:
                    st.error(f'Could not save uploaded files: {e}')

    st.subheader('✍️ Confirm / Edit Technical Inputs')
    col1, col2, col3 = st.columns(3)
    with col1:
        tower_type = st.selectbox('Tower type', options=list(TOWER_TYPES.keys()), format_func=lambda x: TOWER_TYPES[x]['name'], index=list(TOWER_TYPES.keys()).index(extracted.get('tower_type')) if extracted.get('tower_type') in TOWER_TYPES else list(TOWER_TYPES.keys()).index(default_tower_type), key='vendor_tower_type')
        water_flow_m3h = st.number_input('Water flow (m³/h)', min_value=0.0, value=float(extracted.get('water_flow_m3h') or 0.0), step=10.0, format='%.3f')
        hot_water_c = st.number_input('Hot water in (°C)', min_value=-20.0, max_value=120.0, value=float(extracted.get('hot_water_c') or 40.0), step=0.1, format='%.2f')
        cold_water_c = st.number_input('Cold water out (°C)', min_value=-20.0, max_value=120.0, value=float(extracted.get('cold_water_c') or 35.0), step=0.1, format='%.2f')
    with col2:
        wet_bulb_c = st.number_input('Wet bulb (°C)', min_value=-20.0, max_value=80.0, value=float(extracted.get('wet_bulb_c') or 32.0), step=0.1, format='%.2f')
        dry_bulb_c = st.number_input('Dry bulb (°C)', min_value=-20.0, max_value=80.0, value=float(extracted.get('dry_bulb_c') or max((extracted.get('wet_bulb_c') or 32.0) + 3.0, 35.0)), step=0.1, format='%.2f')
        air_flow_m3s = st.number_input('Vendor air flow (m³/s)', min_value=0.0, value=float(extracted.get('air_flow_m3s') or 0.0), step=1.0, format='%.3f')
        fan_power_kw_vendor = st.number_input('Vendor fan power total (kW)', min_value=0.0, value=float(extracted.get('fan_power_kw') or 0.0), step=1.0, format='%.2f')
    with col3:
        static_pressure_pa_vendor = st.number_input('Vendor static pressure (Pa)', min_value=0.0, value=float(extracted.get('static_pressure_pa') or 0.0), step=1.0, format='%.1f')
        length_m = st.number_input('Tower length (m)', min_value=0.0, value=float(extracted.get('length_m') or 0.0), step=0.1, format='%.3f')
        width_m = st.number_input('Tower width (m)', min_value=0.0, value=float(extracted.get('width_m') or 0.0), step=0.1, format='%.3f')
        height_m = st.number_input('Tower height (m)', min_value=0.0, value=float(extracted.get('height_m') or 0.0), step=0.1, format='%.3f')
        fill_depth_m = st.number_input('Comparison fill depth (m)', min_value=0.300, max_value=2.500, value=float(default_fill_depth), step=0.050, format='%.3f')
    face_area = (length_m * width_m) if (length_m and width_m) else 0.0
    if face_area > 0:
        st.info(f'Using face area from vendor dimensions: **{face_area:.2f} m²**')
    else:
        face_area = st.number_input('Fallback face area (m²)', min_value=0.1, value=36.0, step=1.0, format='%.2f')
    l_over_g_guess = st.number_input('Fallback L/G ratio (used if vendor airflow missing)', min_value=0.3, max_value=5.0, value=float(extracted.get('l_over_g') or 1.25), step=0.05, format='%.3f')

    if run_compare:
        errors = []
        if water_flow_m3h <= 0: errors.append('Water flow must be > 0')
        if hot_water_c <= cold_water_c: errors.append('Hot water temperature must be greater than cold water temperature')
        if wet_bulb_c >= hot_water_c: errors.append('Wet bulb should be less than hot water temperature for a sensible comparison')
        if face_area <= 0: errors.append('Face area must be > 0')
        if errors:
            [st.error(e) for e in errors]
            return
        water_mass_kg_s = water_flow_m3h * 1000.0 / 3600.0
        rho_air_vendor = air_density_calc(dry_bulb_c, wet_bulb_c, default_altitude)
        if air_flow_m3s > 0:
            air_mass_kg_s = air_flow_m3s * rho_air_vendor
            l_over_g_used = water_mass_kg_s / air_mass_kg_s if air_mass_kg_s > 0 else None
        else:
            l_over_g_used = l_over_g_guess
            air_mass_kg_s = water_mass_kg_s / l_over_g_used
            air_flow_m3s = air_mass_kg_s / rho_air_vendor if rho_air_vendor > 0 else 0.0
        heat_kw = water_mass_kg_s * 4.186 * (hot_water_c - cold_water_c)
        effective_kavl, effective_ntu = back_calculate_effective_kavl(hot_water_c, cold_water_c, wet_bulb_c, tower_type)
        water_loading = water_flow_m3h / face_area
        air_velocity = air_flow_m3s / face_area
        evap_loss_m3h = estimate_evaporation_loss_m3h(water_flow_m3h, hot_water_c - cold_water_c)
        summary_cols = st.columns(6)
        metrics = [(summary_cols[0],'Heat Rejection',f'{heat_kw:,.0f} kW'),(summary_cols[1],'Effective KaV/L',f'{effective_kavl:.3f}' if effective_kavl is not None else '—'),(summary_cols[2],'Water Loading',f'{water_loading:.1f} m³/h·m²'),(summary_cols[3],'Air Velocity',f'{air_velocity:.2f} m/s'),(summary_cols[4],'L/G',f'{l_over_g_used:.3f}'),(summary_cols[5],'Estimated Evaporation',f'{evap_loss_m3h:.2f} m³/h' if evap_loss_m3h is not None else '—')]
        for c,lbl,val in metrics: c.metric(lbl,val)
        rows = []
        for fill in build_vendor_candidate_fill_list(tower_type):
            use_pdf = (fill=='CT1200AT' and use_pdf_for_ct1200) or (fill=='XF75' and use_pdf_for_xf75)
            res = solve_cooling_tower_enhanced(water_mass_kg_s, air_mass_kg_s, hot_water_c, cold_water_c, wet_bulb_c, dry_bulb_c, fill, tower_type, fill_depth_m, face_area, default_altitude, use_pdf)
            predicted_kavl = res['Ka_over_L'] * fill_depth_m
            rows.append({'Fill': res['fill_name'], 'Data Source':'PDF' if res.get('use_pdf_data', False) else 'Default', 'Vendor Cold Water (°C)': round(cold_water_c,2), 'Model Cold Water (°C)': round(res['T_cold_achieved'],2), 'ΔT Model-Vendor (°C)': round(res['T_cold_achieved']-cold_water_c,2), 'Required Heat (kW)': round(res.get('Q_target', heat_kw),1), 'Achievable Heat (kW)': round(res['Q_achieved'],1), 'Vendor Eff. KaV/L': round(effective_kavl,3) if effective_kavl is not None else None, 'Predicted KaV/L': round(predicted_kavl,3), 'Utilization Ratio': round((effective_kavl/predicted_kavl),3) if effective_kavl is not None and predicted_kavl>0 else None, 'Model Fan Power (kW)': round(res['fan_power'],2), 'Vendor Fan Power (kW)': round(fan_power_kw_vendor,2) if fan_power_kw_vendor else None, 'Water Loading (m³/h·m²)': round(res['water_loading'],2), 'Air Face Velocity (m/s)': round(res['air_face_velocity'],2), 'Net Open Area (m²)': round(res.get('open_air_area',0),2), 'Air Velocity Through Fill (m/s)': round(res.get('air_velocity_through_fill', res['air_face_velocity']),2), 'Free Area Fraction': round(res.get('free_area_fraction',0),2), 'Fill ΔP (Pa)': round(res['fill_pressure_drop'],1), 'Other Losses (Pa)': round(res.get('other_losses_total',0),1), 'Static Pressure (Pa)': round(res['total_static_pressure'],1)})
        df = pd.DataFrame(rows)
        if not df.empty:
            df['ABS ΔT'] = df['ΔT Model-Vendor (°C)'].abs(); df = df.sort_values(['ABS ΔT','Model Fan Power (kW)']).reset_index(drop=True)
        st.subheader('🤖 Automatic Fill Comparison')
        st.dataframe(df, use_container_width=True)
        if not df.empty:
            best = df.iloc[0]
            st.success(f"Best thermal match: **{best['Fill']}** with model cold water **{best['Model Cold Water (°C)']} °C** and ΔT **{best['ΔT Model-Vendor (°C)']} °C**.")
        csv_bytes = df.to_csv(index=False).encode('utf-8')
        st.download_button('📥 Download Vendor Comparison CSV', data=csv_bytes, file_name=f'vendor_fill_comparison_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.csv', mime='text/csv', use_container_width=True)
        vendor_data = {'tower_type': tower_type,'water_flow_m3h': round(water_flow_m3h,3),'hot_water_c': round(hot_water_c,3),'cold_water_c': round(cold_water_c,3),'wet_bulb_c': round(wet_bulb_c,3),'dry_bulb_c': round(dry_bulb_c,3),'length_m': round(length_m,3) if length_m else None,'width_m': round(width_m,3) if width_m else None,'height_m': round(height_m,3) if height_m else None,'face_area_m2': round(face_area,3),'fill_depth_m': round(fill_depth_m,3),'air_flow_m3s': round(air_flow_m3s,3),'air_mass_flow_kg_s': round(air_mass_kg_s,3),'L_over_G': round(l_over_g_used,4),'heat_rejection_kw': round(heat_kw,2),'effective_kavl': round(effective_kavl,4) if effective_kavl is not None else None,'effective_ntu': round(effective_ntu,4) if effective_ntu is not None else None,'estimated_evaporation_m3h': round(evap_loss_m3h,3) if evap_loss_m3h is not None else None,'static_pressure_pa_vendor': round(static_pressure_pa_vendor,2) if static_pressure_pa_vendor else None,'fan_power_kw_vendor': round(fan_power_kw_vendor,2) if fan_power_kw_vendor else None,'source_summary': extracted.get('source_summary', []) if extracted else []}
        report_text = '\n'.join(['='*72,'VENDOR QUOTE COMPARISON REPORT','='*72,f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",'','EXTRACTED / CONFIRMED VENDOR DATA','-'*72]+[f"{k}: {v}" for k,v in vendor_data.items() if k!='source_summary']+['','EXTRACTION NOTES','-'*72]+[f"- {x}" for x in vendor_data.get('source_summary',[])] + ['','FILL COMPARISON','-'*72, df.to_string(index=False), '', '='*72])
        st.download_button('📥 Download Vendor Comparison TXT', data=report_text, file_name=f'vendor_fill_comparison_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.txt', mime='text/plain', use_container_width=True)
        pdf_bytes = generate_pdf_report_from_text(report_text, title='Vendor Quote Comparison Report')
        st.download_button('📥 Download Vendor Comparison PDF', data=pdf_bytes, file_name=f'vendor_fill_comparison_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf', mime='application/pdf', use_container_width=True)
        docx_bytes = generate_docx_report('Vendor Quote Comparison Report', [('Confirmed Vendor Inputs', vendor_data), ('Fill Comparison', df)])
        st.download_button('📥 Download Vendor Comparison Word (.docx)', data=docx_bytes, file_name=f'vendor_fill_comparison_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', use_container_width=True)
        save_row = {'saved_at': datetime.datetime.now().isoformat(timespec='seconds'), 'source_files': ', '.join([uf.name for uf in (uploaded_files or [])]), 'tower_type': tower_type, 'water_flow_m3h': water_flow_m3h, 'hot_water_c': hot_water_c, 'cold_water_c': cold_water_c, 'wet_bulb_c': wet_bulb_c, 'dry_bulb_c': dry_bulb_c, 'length_m': length_m, 'width_m': width_m, 'height_m': height_m if height_m else None, 'heat_load_kw': heat_kw, 'effective_kavl': effective_kavl, 'air_flow_m3s': air_flow_m3s, 'vendor_fan_power_kw': fan_power_kw_vendor, 'vendor_static_pressure_pa': static_pressure_pa_vendor}
        if auto_save_to_library:
            ok = save_spec_to_library(save_row)
            st.success('Parsed spec saved to library.' if ok else 'Could not save parsed spec to library.')
        if google_drive_is_configured() and auto_save_originals:
            cached_vendor_files = get_cached_vendor_files()
            if cached_vendor_files:
                try:
                    saved, errors = save_cached_vendor_files_to_drive()
                    if saved:
                        st.success('Original uploaded vendor files saved to Google Drive.')
                    if errors:
                        st.error('Some original files could not be saved: ' + ' | '.join(errors))
                except Exception as e:
                    st.error(f'Could not save original vendor files: {e}')
        if google_drive_is_configured():
            try:
                stamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
                upload_json_to_drive(f'vendor_summary_{stamp}.json', vendor_data)
                upload_text_to_drive(f'vendor_fill_comparison_{stamp}.txt', report_text)
            except Exception as e:
                st.warning(f'Could not upload summary/report to Google Drive: {e}')
        with st.expander('Preview vendor comparison report'):
            st.text(report_text[:4000] + ('...' if len(report_text) > 4000 else ''))


if __name__ == "__main__":
    main()